"""
shira_proxy.py — Local proxy server for Shira API
Requirements:
    pip install flask requests requests-negotiate-sspi beautifulsoup4 lxml python-docx pdfplumber flask-cors httpx
"""

import os
import ssl
import urllib3
import re
import io
import json
import xml.etree.ElementTree as ET

# ── Kill any existing process on port 5050 before starting ───────────────────
def _kill_port(port=5050):
    import subprocess, signal, sys
    try:
        if sys.platform == "win32":
            result = subprocess.check_output(
                f'netstat -ano | findstr :{port}', shell=True
            ).decode(errors='ignore')
            pids = set()
            for line in result.splitlines():
                parts = line.split()
                if len(parts) >= 5 and 'LISTENING' in line:
                    pids.add(parts[-1])
            for pid in pids:
                try:
                    subprocess.call(f'taskkill /PID {pid} /F', shell=True,
                                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    print(f'[startup] closed old server PID {pid}')
                except Exception:
                    pass
    except Exception:
        pass

_kill_port(5050)

os.environ['NO_PROXY'] = 'shira2,prod-spfe,10.67.60.51,localhost,127.0.0.1'
urllib3.disable_warnings()
ssl._create_default_https_context = ssl._create_unverified_context

from flask import Flask, request, jsonify, Response, stream_with_context
from flask_cors import CORS
import requests
from requests_negotiate_sspi import HttpNegotiateAuth
from bs4 import BeautifulSoup
import pdfplumber
import docx

app = Flask(__name__)
app.config['JSON_AS_ASCII'] = False
CORS(app)

# ── Embedded frontend HTML ────────────────────────────────────────────────────
# The HTML is embedded here so the EXE is fully self-contained.
_HTML = """<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>מערכת שירה</title>
<style>
* { box-sizing: border-box; margin: 0; padding: 0; }
body {
  font-family: 'Segoe UI', Arial, sans-serif;
  direction: rtl;
  background: #f4f5f7;
  color: #1a1a2e;
  font-size: 14px;
}
header {
  background: #1a3a5c;
  color: #fff;
  padding: 12px 24px;
  display: flex;
  align-items: center;
  gap: 12px;
}
header h1 { font-size: 17px; font-weight: 500; }
header .sub { font-size: 12px; opacity: 0.7; margin-top: 2px; }
.status-dot {
  width: 10px; height: 10px; border-radius: 50%;
  background: #ccc; margin-right: auto; margin-left: 8px;
  transition: background 0.3s;
}
.status-dot.ok  { background: #4caf50; }
.status-dot.err { background: #f44336; }
.status-label { font-size: 12px; opacity: 0.8; }
.user-chip {
  font-size: 12px;
  background: rgba(255,255,255,0.15);
  border-radius: 20px;
  padding: 4px 12px;
  opacity: 0.9;
}
.container { max-width: 1100px; margin: 0 auto; padding: 20px 16px; }
.card {
  background: #fff;
  border: 1px solid #e0e4ea;
  border-radius: 10px;
  padding: 20px;
  margin-bottom: 16px;
}
.card-title {
  font-size: 13px; font-weight: 600; color: #555;
  margin-bottom: 14px; display: flex; align-items: center; gap: 8px;
  text-transform: uppercase; letter-spacing: 0.5px;
}
.row { display: flex; gap: 8px; align-items: center; margin-bottom: 12px; }
input[type=text] {
  flex: 1; height: 38px;
  border: 1px solid #d0d5dd; border-radius: 7px;
  padding: 0 12px; font-size: 14px; direction: rtl;
  outline: none; transition: border 0.2s;
}
input[type=text]:focus { border-color: #1a3a5c; box-shadow: 0 0 0 3px rgba(26,58,92,0.12); }
input[type=text]::placeholder { color: #aaa; }
button {
  height: 38px; padding: 0 18px;
  border: 1px solid #d0d5dd; border-radius: 7px;
  background: #fff; color: #1a1a2e;
  font-size: 14px; cursor: pointer;
  transition: all 0.15s; white-space: nowrap;
}
button:hover  { background: #f0f4f8; }
button:active { transform: scale(0.98); }
button.primary { background: #1a3a5c; color: #fff; border-color: #1a3a5c; }
button.primary:hover { background: #14304d; }
button.sm { height: 30px; padding: 0 12px; font-size: 12px; }
table { width: 100%; border-collapse: collapse; font-size: 13px; }
th {
  text-align: right; padding: 9px 12px;
  background: #f8f9fb; color: #555; font-weight: 600;
  border-bottom: 1px solid #e0e4ea; position: sticky; top: 0;
}
td { padding: 9px 12px; border-bottom: 1px solid #f0f2f5; vertical-align: middle; }
tr.clickable:hover { background: #f0f6ff; cursor: pointer; }
tr.selected { background: #e8f0fe !important; }
.badge { display: inline-block; padding: 2px 9px; border-radius: 20px; font-size: 11px; font-weight: 600; }
.b-open    { background: #e6f4ea; color: #2e7d32; }
.b-closed  { background: #f3f4f6; color: #555; }
.b-pending { background: #fff8e1; color: #e65100; }
.tabs { display: flex; gap: 4px; margin-bottom: 16px; border-bottom: 1px solid #e0e4ea; }
.tab {
  padding: 8px 16px; border-radius: 7px 7px 0 0;
  font-size: 13px; cursor: pointer;
  border: 1px solid transparent; background: transparent;
  color: #666; border-bottom: none; position: relative; bottom: -1px;
}
.tab.active { background: #fff; border-color: #e0e4ea; border-bottom-color: #fff; color: #1a3a5c; font-weight: 600; }
.tab:hover:not(.active) { background: #f0f4f8; }
.empty { color: #aaa; text-align: center; padding: 28px; font-size: 13px; }
.spinner {
  display: inline-block; width: 16px; height: 16px;
  border: 2px solid #ddd; border-top-color: #1a3a5c;
  border-radius: 50%; animation: spin 0.7s linear infinite;
  vertical-align: middle; margin-left: 6px;
}
@keyframes spin { to { transform: rotate(360deg); } }
.ai-box {
  background: #f8f9fb; border: 1px solid #e0e4ea;
  border-radius: 8px; padding: 14px 16px;
  font-size: 13px; line-height: 1.75;
  white-space: pre-wrap; min-height: 80px; color: #222;
}
.ai-box.loading { color: #999; font-style: italic; }
.ai-cursor { display: inline-block; width: 2px; height: 14px; background: #1a3a5c; animation: blink 0.8s step-end infinite; vertical-align: middle; margin-right: 2px; }
@keyframes blink { 50% { opacity: 0; } }
.doc-row {
  display: flex; align-items: center; gap: 10px;
  padding: 9px 4px; border-bottom: 1px solid #f0f2f5;
}
.doc-row:hover { background: #f8f9fb; }
.doc-name { flex: 1; font-size: 13px; }
.doc-date { color: #888; font-size: 12px; white-space: nowrap; }
.doc-type { font-size: 11px; color: #888; background: #f0f2f5; padding: 1px 7px; border-radius: 10px; }
.match-badge { font-size: 11px; color: #e65100; background: #fff3e0; padding: 2px 7px; border-radius: 10px; }
mark { background: #fff176; border-radius: 2px; padding: 0 1px; }
.stat-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 10px; margin-bottom: 16px; }
.stat { background: #f8f9fb; border-radius: 8px; padding: 12px 14px; text-align: center; }
.stat-val { font-size: 20px; font-weight: 600; color: #1a3a5c; }
.stat-lbl { font-size: 11px; color: #888; margin-top: 3px; }
.hearing-row { padding: 8px 4px; border-bottom: 1px solid #f0f2f5; font-size: 13px; display: flex; gap: 12px; align-items: center; }
.hearing-date { font-weight: 600; color: #1a3a5c; min-width: 90px; }
.soon { background: #fff3e0; color: #e65100; padding: 1px 7px; border-radius: 10px; font-size: 11px; }
.error-msg { color: #c62828; background: #ffebee; border-radius: 6px; padding: 8px 12px; font-size: 13px; }
.mode-banner { border-radius: 8px; padding: 8px 14px; font-size: 12px; margin-bottom: 16px; font-weight: 500; }
.mode-banner-user { background: #e8f5e9; color: #2e7d32; border: 1px solid #c8e6c9; }
.mode-banner-dev  { background: #fff8e1; color: #e65100; border: 1px solid #ffe0b2; }
</style>
</head>
<body>

<header>
  <div>
    <h1 id="header-title">מערכת שירה — חיפוש חכם וסיכומי AI</h1>
    <div class="sub">בתי הדין הרבניים</div>
  </div>
  <div style="margin-right:auto;display:flex;align-items:center;gap:10px;">
    <span class="user-chip" id="user-chip"></span>
  </div>
  <div class="status-dot" id="dot"></div>
  <div class="status-label" id="status-label">מתחבר...</div>
</header>

<button id="dev-btn" onclick="toggleDevMode()" style="position:fixed;bottom:16px;left:16px;z-index:999;background:rgba(0,0,0,0.06);border:1px solid rgba(0,0,0,0.1);color:#aaa;font-size:11px;padding:5px 10px;border-radius:20px;cursor:pointer;height:auto;opacity:0.4;transition:opacity 0.2s;" onmouseenter="this.style.opacity=1" onmouseleave="this.style.opacity=0.4">⚙</button>

<button id="usage-btn" onclick="showUsage()" style="display:none;position:fixed;bottom:16px;left:80px;z-index:999;background:rgba(0,0,0,0.06);border:1px solid rgba(0,0,0,0.1);color:#aaa;font-size:11px;padding:5px 10px;border-radius:20px;cursor:pointer;height:auto;opacity:0.4;transition:opacity 0.2s;" onmouseenter="this.style.opacity=1" onmouseleave="this.style.opacity=0.4">📊</button>

<!-- Usage popup -->
<div id="usage-popup" style="display:none;position:fixed;bottom:50px;left:16px;z-index:1000;background:#fff;border:1px solid #e0e4ea;border-radius:12px;padding:20px;min-width:320px;box-shadow:0 4px 24px rgba(0,0,0,0.12);">
  <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:14px;">
    <strong style="font-size:13px;color:#1a3a5c;">📊 סטטיסטיקות שימוש</strong>
    <button onclick="document.getElementById('usage-popup').style.display='none'" style="height:24px;padding:0 8px;font-size:12px;border-radius:6px;">✕</button>
  </div>
  <div id="usage-content" style="font-size:13px;color:#444;line-height:2;">טוען...</div>
</div>

<div class="container">

  <div id="mode-banner" style="display:none"></div>

  <div class="card">
    <div class="card-title">🔍 חיפוש תיקים</div>
    <div style="display:flex;gap:4px;margin-bottom:14px;border-bottom:1px solid #e0e4ea;padding-bottom:0;">
      <button class="tab active" id="stab-id"   onclick="switchSearchTab('id')">לפי ת"ז</button>
      <button class="tab"        id="stab-case" onclick="switchSearchTab('case')">לפי מס' תיק</button>
      <button class="tab"        id="stab-name" onclick="switchSearchTab('name')">לפי שם</button>
    </div>
    <div id="search-id-panel">
      <div class="row">
        <input type="text" id="id-input" placeholder='הכנס מספר ת"ז (9 ספרות)' maxlength="9" />
        <button class="primary" onclick="doSearch()">חפש</button>
      </div>
    </div>
    <div id="search-case-panel" style="display:none">
      <div class="row">
        <input type="text" id="case-input" placeholder='מס׳ תיק (לדוגמה: 1488524 או 1488524/1)' />
        <button class="primary" onclick="doCaseSearch()">חפש</button>
      </div>
    </div>
    <div id="search-name-panel" style="display:none">
      <div class="row">
        <input type="text" id="name-first" placeholder="שם פרטי" />
        <input type="text" id="name-last"  placeholder="שם משפחה" />
        <button class="primary" onclick="doNameSearch()">חפש</button>
      </div>
    </div>
    <div id="results-area"></div>
  </div>

  <div class="card" id="case-panel" style="display:none">
    <div class="card-title" id="case-heading">📁 תיק נבחר</div>
    <div class="stat-grid" id="case-stats"></div>
    <div class="tabs">
      <button class="tab active" onclick="switchTab('docs')">📄 מסמכים</button>
      <button class="tab" onclick="switchTab('hearings')">📅 דיונים</button>
      <button class="tab" onclick="switchTab('search')">🔎 חיפוש בתיק</button>
      <button class="tab" onclick="switchTab('ai')">✨ סיכום AI</button>
      <button class="tab" onclick="switchTab('msg')">📨 שלח הודעה</button>
    </div>
    <div id="tab-docs"></div>
    <div id="tab-hearings" style="display:none"></div>
    <div id="tab-search"   style="display:none"></div>
    <div id="tab-ai"       style="display:none"></div>
    <div id="tab-msg"      style="display:none"></div>
  </div>

</div>

<script>
const PROXY = "http://localhost:5050";

const COURT_NAMES = {
  1:'ירושלים', 2:'תל אביב', 3:'חיפה', 4:'פתח תקוה',
  5:'רחובות',  6:'באר שבע', 7:'טבריה', 8:'צפת',
  9:'אשדוד',  10:'אשקלון', 11:'נתניה',
  12:'בית הדין הגדול', 13:'אריאל'
};

// State
let userCourtId   = null;   // set from /api/me
let userCourtName = null;
let userName      = null;
let devMode       = false;
let caseDocs      = [];
let docTexts      = {};
let selectedCase  = null;

// ── Boot ──────────────────────────────────────────────────────────────────────
async function boot() {
  checkHealth();
  setInterval(checkHealth, 30000);
  try {
    const r = await fetch(`${PROXY}/api/me`);
    const d = await r.json();
    if (d.courtId) {
      userCourtId   = d.courtId;
      userCourtName = d.courtName || COURT_NAMES[d.courtId] || String(d.courtId);
      userName      = d.firstName ? `${d.firstName} ${d.lastName}` : d.userName || '';
      document.getElementById('header-title').textContent =
        `חיפוש חכם וסיכומי AI — ${userCourtName}`;
      document.title = `חיפוש חכם וסיכומי AI — ${userCourtName}`;
      document.getElementById('user-chip').textContent = userName ? `👤 ${userName}` : '';
      setBanner(false);
    } else {
      document.getElementById('mode-banner').textContent = '⚠️ לא ניתן לזהות בית דין — פנה למנהל המערכת';
    }
  } catch(e) {
    document.getElementById('mode-banner').textContent = '⚠️ שגיאה בזיהוי בית דין: ' + e.message;
  }
}

function setBanner(dev) {
  const btn     = document.getElementById('dev-btn');
  const usageBtn = document.getElementById('usage-btn');
  const costEls = document.querySelectorAll('.ai-cost');
  if (dev) {
    btn.textContent = '⚙ מצב מפתח';
    btn.style.color = '#e65100';
    btn.style.borderColor = 'rgba(230,81,0,0.3)';
    usageBtn.style.display = 'block';
    costEls.forEach(el => el.style.display = 'inline');
  } else {
    btn.textContent = '⚙';
    btn.style.color = '#aaa';
    btn.style.borderColor = 'rgba(0,0,0,0.1)';
    usageBtn.style.display = 'none';
    document.getElementById('usage-popup').style.display = 'none';
    costEls.forEach(el => el.style.display = 'none');
  }
}

function toggleDevMode() {
  if (!devMode) {
    const pwd = prompt('סיסמת מפתח:');
    if (pwd !== 'ELCH2026') {
      alert('סיסמה שגויה');
      return;
    }
  }
  devMode = !devMode;
  setBanner(devMode);
  document.getElementById('results-area').innerHTML = '';
  document.getElementById('case-panel').style.display = 'none';
}

// ── Health ────────────────────────────────────────────────────────────────────
async function checkHealth() {
  try {
    const r = await fetch(`${PROXY}/api/health`);
    document.getElementById('dot').className = r.ok ? 'status-dot ok' : 'status-dot err';
    document.getElementById('status-label').textContent = r.ok ? 'מחובר לשרת' : 'שגיאה';
  } catch {
    document.getElementById('dot').className = 'status-dot err';
    document.getElementById('status-label').textContent = 'שגיאה: הפעל את shira_proxy.py';
  }
}

// ── Court filter ──────────────────────────────────────────────────────────────
function applyCourtFilter(data) {
  if (devMode || !userCourtId) return data;
  return data.filter(c => c.courtId === userCourtId);
}

// ── Search tab switch ─────────────────────────────────────────────────────────
function switchSearchTab(tab) {
  document.getElementById('search-id-panel').style.display   = tab === 'id'   ? 'block' : 'none';
  document.getElementById('search-case-panel').style.display = tab === 'case' ? 'block' : 'none';
  document.getElementById('search-name-panel').style.display = tab === 'name' ? 'block' : 'none';
  document.getElementById('stab-id').classList.toggle('active',   tab === 'id');
  document.getElementById('stab-case').classList.toggle('active', tab === 'case');
  document.getElementById('stab-name').classList.toggle('active', tab === 'name');
  document.getElementById('results-area').innerHTML = '';
}

// ── Search by name ────────────────────────────────────────────────────────────
async function doNameSearch() {
  const firstName = document.getElementById('name-first').value.trim();
  const lastName  = document.getElementById('name-last').value.trim();
  const area      = document.getElementById('results-area');
  if (!lastName && !firstName) { area.innerHTML = '<p class="empty">נא להכניס שם משפחה או שם פרטי</p>'; return; }
  area.innerHTML = '<p class="empty"><span class="spinner"></span> מחפש...</p>';
  document.getElementById('case-panel').style.display = 'none';
  try {
    const r = await fetch(`${PROXY}/api/search-name`, {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify({ lastName, firstName })
    });
    const data = await r.json();
    if (data.error) { area.innerHTML = `<div class="error-msg">שגיאה: ${data.error}</div>`; return; }
    renderResults(Array.isArray(data) ? data : [], area);
  } catch {
    area.innerHTML = `<div class="error-msg">לא ניתן להתחבר לשרת. ודא ש-shira_proxy.py פועל.</div>`;
  }
}

// ── Render results ────────────────────────────────────────────────────────────
function renderResults(data, area) {
  const filtered = applyCourtFilter(data);
  if (!filtered.length) {
    area.innerHTML = `<p class="empty">${!devMode && data.length > 0 ? 'לא נמצאו תיקים בבית הדין שלך' : 'לא נמצאו תיקים'}</p>`;
    return;
  }
  window._cases = filtered;
  let html = `
    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:8px;">
      <p style="font-size:12px;color:#888;">נמצאו ${filtered.length} תיקים</p>
      <button class="sm primary" onclick="showSearchAllPanel()" style="font-size:12px;">🔎 חפש בכל התיקים</button>
    </div>
    <div id="search-all-panel" style="display:none;background:#f8f9fb;border:1px solid #e0e4ea;border-radius:8px;padding:12px;margin-bottom:12px;">
      <p style="font-size:12px;color:#555;margin-bottom:8px;">חיפוש בתוכן מסמכי כל ${filtered.length} התיקים</p>
      <div class="row">
        <input type="text" id="search-all-q" placeholder="הכנס מילה או ביטוי לחיפוש..." />
        <button class="primary" onclick="doSearchAllCases()">חפש</button>
      </div>
      <div id="search-all-results"></div>
    </div>
    <table><thead><tr>
      <th>מס' תיק</th><th>בית דין</th><th>נושא</th><th>צד א</th><th>צד ב</th><th>סטטוס</th>
    </tr></thead><tbody>`;
  filtered.forEach((c, i) => {
    const cls = c.fileStatusID === 2 ? 'b-closed' : c.fileStatusID === 5 ? 'b-pending' : 'b-open';
    const status = c.fileStatusName || (c.isClosed ? 'סגור' : 'פתוח');
    html += `<tr class="clickable" onclick="selectCase(${i})" data-idx="${i}">
      <td>${c.fullFileMainNumber || c.fileNumber || ''}</td>
      <td>${c.courtName || ''}</td>
      <td>${c.subjectSubName || ''}</td>
      <td>${(c.sideA || '').substring(0,25)}</td>
      <td>${(c.sideB || '').substring(0,25)}</td>
      <td><span class="badge ${cls}">${status}</span></td>
    </tr>`;
  });
  html += '</tbody></table>';
  area.innerHTML = html;
}

function showSearchAllPanel() {
  const panel = document.getElementById('search-all-panel');
  if (!panel) return;
  panel.style.display = panel.style.display === 'none' ? 'block' : 'none';
  if (panel.style.display === 'block') {
    document.getElementById('search-all-q').focus();
  }
}

async function doSearchAllCases() {
  const q    = document.getElementById('search-all-q').value.trim();
  const area = document.getElementById('search-all-results');
  if (!q) return;

  const cases = window._cases || [];
  if (!cases.length) { area.innerHTML = '<p class="empty">אין תיקים לחיפוש</p>'; return; }

  area.innerHTML = `<p class="empty"><span class="spinner"></span> טוען מסמכים מ-${cases.length} תיקים...</p>`;

  let totalHits = 0;
  let html = '';

  for (const c of cases) {
    const fileId = c.fileId || c.fileMainId;

    // Load docs
    if (!window._allCaseDocs) window._allCaseDocs = {};
    if (!window._allCaseDocs[fileId]) {
      try {
        const r = await fetch(`${PROXY}/api/documents/${fileId}`);
        window._allCaseDocs[fileId] = await r.json();
      } catch { window._allCaseDocs[fileId] = []; }
    }
    const docs = window._allCaseDocs[fileId] || [];

    // Load texts and search
    const hits = [];
    for (const doc of docs) {
      const key = `all_${doc.docId}`;
      if (!docTexts[key]) {
        try {
          const r = await fetch(`${PROXY}/api/doctext/${doc.docId}`);
          const d = await r.json();
          docTexts[key] = d.text || '';
        } catch { docTexts[key] = ''; }
      }
      if (doc.name.includes(q) || (docTexts[key]||'').includes(q)) {
        hits.push({ ...doc, _textKey: key });
      }
    }

    area.innerHTML = `<p class="empty"><span class="spinner"></span> בודק תיק ${c.fileNumber}...</p>`;

    if (hits.length) {
      totalHits += hits.length;
      html += `<div style="margin-bottom:14px;">
        <div style="font-size:12px;font-weight:600;color:#1a3a5c;padding:6px 0;border-bottom:1px solid #e0e4ea;margin-bottom:8px;">
          📁 ${c.fileNumber} — ${c.subjectSubName || ''}
          <span style="font-weight:400;color:#888;margin-right:8px;">${hits.length} תוצאות</span>
        </div>`;
      hits.forEach(doc => {
        const text    = docTexts[doc._textKey] || '';
        const idx     = text.indexOf(q);
        const snippet = idx >= 0
          ? '...' + text.substring(Math.max(0,idx-40), idx+q.length+80)
              .replace(new RegExp(q,'g'), `<mark>${q}</mark>`) + '...'
          : '';
        html += `<div class="doc-row" style="flex-direction:column;align-items:flex-start;gap:4px;">
          <div style="display:flex;gap:8px;align-items:center;width:100%">
            <span>${doc.type==='pdf'?'📕':'📄'}</span>
            <strong style="font-size:13px">${doc.name}</strong>
            <span class="doc-date">${doc.date}</span>
            <button class="sm primary" style="margin-right:auto" onclick="openDoc('${encodeURIComponent(doc.openUrl)}')">פתח</button>
          </div>
          ${snippet ? `<p style="font-size:12px;color:#555;padding-right:24px;line-height:1.7">${snippet}</p>` : ''}
        </div>`;
      });
      html += '</div>';
    }
  }

  if (!totalHits) {
    area.innerHTML = `<p class="empty">לא נמצאו תוצאות עבור "${q}" ב-${cases.length} תיקים</p>`;
    return;
  }
  area.innerHTML = `<p style="font-size:12px;color:#888;margin-bottom:12px;">נמצאו ${totalHits} מסמכים ב-${cases.length} תיקים</p>` + html;
}

// ── Search by ID ──────────────────────────────────────────────────────────────
async function doSearch() {
  const idNum = document.getElementById('id-input').value.trim();
  const area  = document.getElementById('results-area');
  if (!idNum || idNum.length < 5) { area.innerHTML = '<p class="empty">נא להכניס ת"ז תקינה</p>'; return; }
  area.innerHTML = '<p class="empty"><span class="spinner"></span> מחפש...</p>';
  document.getElementById('case-panel').style.display = 'none';
  try {
    const r = await fetch(`${PROXY}/api/search`, {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify({ idNum })
    });
    const data = await r.json();
    if (data.error) { area.innerHTML = `<div class="error-msg">שגיאה: ${data.error}</div>`; return; }
    renderResults(Array.isArray(data) ? data : [], area);
  } catch {
    area.innerHTML = `<div class="error-msg">לא ניתן להתחבר לשרת. ודא ש-shira_proxy.py פועל.</div>`;
  }
}

// ── Search by case number ─────────────────────────────────────────────────────
async function doCaseSearch() {
  const raw  = document.getElementById('case-input').value.trim();
  const area = document.getElementById('results-area');
  if (!raw) { area.innerHTML = '<p class="empty">נא להכניס מספר תיק</p>'; return; }
  const parts      = raw.split('/');
  const fileMainId = parts[0].replace(/\\D/g, '');
  const fileNumber = parts[1] ? parts[1].replace(/\\D/g, '') : null;
  if (!fileMainId) { area.innerHTML = '<p class="empty">מספר תיק לא תקין</p>'; return; }
  area.innerHTML = '<p class="empty"><span class="spinner"></span> מחפש...</p>';
  document.getElementById('case-panel').style.display = 'none';
  try {
    const r = await fetch(`${PROXY}/api/search-case`, {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify({ fileMainId, fileNumber })
    });
    const data = await r.json();
    if (data.error) { area.innerHTML = `<div class="error-msg">שגיאה: ${data.error}</div>`; return; }
    renderResults(Array.isArray(data) ? data : (data ? [data] : []), area);
  } catch {
    area.innerHTML = `<div class="error-msg">לא ניתן להתחבר לשרת. ודא ש-shira_proxy.py פועל.</div>`;
  }
}

// ── Select case ───────────────────────────────────────────────────────────────
async function selectCase(idx) {
  selectedCase = window._cases[idx];
  caseDocs = []; docTexts = {}; _preloadAbort = true;
  document.getElementById('tab-ai').innerHTML = '';  // clear so renderAITab re-runs for new case
  document.querySelectorAll('tr[data-idx]').forEach(r => r.classList.remove('selected'));
  document.querySelector(`tr[data-idx="${idx}"]`)?.classList.add('selected');
  document.getElementById('case-heading').textContent =
    `📁 תיק ${selectedCase.fullFileMainNumber || selectedCase.fileNumber} — ${selectedCase.sideA || ''} / ${selectedCase.sideB || ''}`;
  document.getElementById('case-stats').innerHTML = `
    <div class="stat"><div class="stat-val">${selectedCase.courtName||'—'}</div><div class="stat-lbl">בית דין</div></div>
    <div class="stat"><div class="stat-val">${selectedCase.subjectSubName||'—'}</div><div class="stat-lbl">נושא</div></div>
    <div class="stat"><div class="stat-val">${selectedCase.isClosed?'סגור':'פתוח'}</div><div class="stat-lbl">סטטוס</div></div>
  `;
  document.getElementById('case-panel').style.display = 'block';
  switchTab('docs');
  document.getElementById('case-panel').scrollIntoView({ behavior: 'smooth' });
  loadDocs();
}

// ── Load documents ────────────────────────────────────────────────────────────
async function loadDocs() {
  const fileId = selectedCase.fileId || selectedCase.fileMainId;
  try {
    const r = await fetch(`${PROXY}/api/documents/${fileId}`);
    const data = await r.json();
    caseDocs = data.error ? [] : data;
    if (document.getElementById('tab-docs').style.display !== 'none') renderDocs('');
    // Start background preloading after docs list is ready
    preloadDocTexts();
  } catch { caseDocs = []; }
}

// ── Background preloading ─────────────────────────────────────────────────────
let _preloadAbort = false;
async function preloadDocTexts() {
  _preloadAbort = false;
  const toLoad = caseDocs.filter(d => !docTexts[d.docId]);
  if (!toLoad.length) return;
  console.log(`[preload] starting background load of ${toLoad.length} docs`);
  for (const doc of toLoad) {
    if (_preloadAbort) break;
    if (docTexts[doc.docId]) continue;
    try {
      const ctrl = new AbortController();
      const tid = setTimeout(() => ctrl.abort(), 8000);
      try {
        const r = await fetch(`${PROXY}/api/doctext/${doc.docId}`, { signal: ctrl.signal });
        clearTimeout(tid);
        const d = await r.json();
        docTexts[doc.docId] = d.text || '';
      } catch { clearTimeout(tid); docTexts[doc.docId] = ''; }
    } catch { docTexts[doc.docId] = ''; }
    // Small delay to avoid overwhelming the server
    await new Promise(res => setTimeout(res, 150));
  }
  console.log(`[preload] done — ${Object.keys(docTexts).length} docs in memory`);
  renderAIDocList(); // refresh checkmarks
}

// ── Tabs ──────────────────────────────────────────────────────────────────────
function switchTab(tab) {
  document.querySelectorAll('.tab').forEach((t, i) =>
    t.classList.toggle('active', ['docs','hearings','search','ai','msg'][i] === tab));
  ['docs','hearings','search','ai','msg'].forEach(t => {
    document.getElementById('tab-'+t).style.display = t === tab ? 'block' : 'none';
  });
  if (tab === 'docs')     renderDocs('');
  if (tab === 'hearings') loadHearings();
  if (tab === 'search')   renderSearchTab();
  if (tab === 'ai')       renderAITab();
  if (tab === 'msg')      renderMsgTab();
}

// ── Message templates ─────────────────────────────────────────────────────────
const MSG_TEMPLATES = [
  { id: 'zimun',    name: 'זימון לדיון',           text: 'הנכם מוזמנים להתייצב לדיון בתיק הנ"ל אשר יתקיים במועד שייקבע ויודע לכם בנפרד.\nהנכם מתבקשים לאשר קבלת הודעה זו.' },
  { id: 'dachuy',   name: 'דחיית דיון',            text: 'הדיון שנקבע בתיק הנ"ל נדחה למועד חדש אשר יודע לכם בנפרד.\nנא לאשר קבלת הודעה זו.' },
  { id: 'mesamchim',name: 'בקשה להגשת מסמכים',    text: 'הנכם מתבקשים להגיש לבית הדין את המסמכים הרלוונטיים לתיק הנ"ל בתוך 14 יום מקבלת הודעה זו.' },
  { id: 'hachlatah',name: 'הודעה על החלטה',        text: 'בית הדין מודיע כי ניתנה החלטה בתיק הנ"ל. ההחלטה מצורפת בזה.' },
  { id: 'hofshi',   name: 'נוסח חופשי',            text: '' },
];

function renderMsgTab() {
  const el = document.getElementById('tab-msg');
  const stored = JSON.parse(localStorage.getItem('shira_msg_templates') || '[]');
  const allTemplates = [...MSG_TEMPLATES, ...stored];

  const btnStyle = 'padding:6px 12px;margin:3px;border:1px solid #b0bec5;border-radius:6px;cursor:pointer;font-size:13px;background:#f5f5f5';
  const activeBtnStyle = btnStyle.replace('#f5f5f5','#1a3a5c').replace('border:1px solid #b0bec5','border:1px solid #1a3a5c') + ';color:#fff';

  const templateBtns = allTemplates.map(function(t, i) {
    const st = i === 0 ? activeBtnStyle : btnStyle;
    return '<button style="' + st + '" onclick="selectMsgTemplate(' + i + ')">' + t.name + '</button>';
  }).join('');

  el.innerHTML =
    '<div style="padding:16px 0">' +
    '<div style="margin-bottom:10px;font-weight:600;font-size:14px;color:#1a3a5c">בחר תבנית:</div>' +
    '<div id="msg-template-btns">' + templateBtns +
    '<button style="' + btnStyle + '" onclick="saveMsgTemplate()">＋ שמור תבנית חדשה</button></div>' +
    '<div style="margin-top:14px;margin-bottom:6px;font-weight:600;font-size:14px;color:#1a3a5c">תוכן ההודעה:</div>' +
    '<textarea id="msg-text" dir="rtl" style="width:100%;height:160px;padding:10px;font-size:14px;font-family:FrankRuehl,Arial;border:1px solid #b0bec5;border-radius:6px;resize:vertical;box-sizing:border-box"></textarea>' +
    '<div style="margin-top:12px">' +
    '<div style="margin-bottom:8px;font-weight:600;font-size:14px;color:#1a3a5c">אופן שליחה:</div>' +
    '<label style="margin-left:16px;cursor:pointer"><input type="radio" name="msg-send-mode" value="a" checked onchange="msgModeChanged()"> <b>א — מתוך התיק</b> (נוצר מסמך בתיק, נשלח דרך מסך הדיוור של שירה)</label><br>' +
    '<label style="cursor:pointer"><input type="radio" name="msg-send-mode" value="b" onchange="msgModeChanged()"> <b>ב — ישיר ללא מסמך</b> (אימייל מ-no-reply@rbc.gov.il, לא נשמר בתיק)</label>' +
    '</div>' +
    '<div id="msg-email-row" style="display:none;margin-top:10px">' +
    '<label style="font-size:13px;font-weight:600;color:#1a3a5c">כתובת אימייל של הנמען:</label><br>' +
    '<input id="msg-to-email" type="email" dir="ltr" placeholder="example@domain.com" style="width:100%;padding:8px;margin-top:4px;border:1px solid #b0bec5;border-radius:6px;font-size:14px;box-sizing:border-box">' +
    '</div>' +
    '<div style="margin-top:12px;text-align:center">' +
    '<button id="msg-send-btn" onclick="sendMessage()" style="padding:10px 32px;background:#1a3a5c;color:#fff;border:none;border-radius:8px;font-size:15px;cursor:pointer;font-weight:600">📨 שלח הודעה</button>' +
    '</div><div id="msg-status" style="margin-top:10px;text-align:center;font-size:13px"></div></div>';

  document.getElementById('msg-text').value = allTemplates[0].text;
  window._msgTemplates = allTemplates;
  window._msgSelectedIdx = 0;
}

function selectMsgTemplate(idx) {
  window._msgSelectedIdx = idx;
  const t = window._msgTemplates[idx];
  document.getElementById('msg-text').value = t.text;
  const btns = document.querySelectorAll('#msg-template-btns button');
  const base = 'padding:6px 12px;margin:3px;border:1px solid #b0bec5;border-radius:6px;cursor:pointer;font-size:13px;background:#f5f5f5';
  const active = 'padding:6px 12px;margin:3px;border:1px solid #1a3a5c;border-radius:6px;cursor:pointer;font-size:13px;background:#1a3a5c;color:#fff';
  btns.forEach((b, i) => { if (i < window._msgTemplates.length) b.style.cssText = (i === idx ? active : base); });
}

function saveMsgTemplate() {
  const name = prompt('שם התבנית החדשה:');
  if (!name) return;
  const text = document.getElementById('msg-text').value;
  const stored = JSON.parse(localStorage.getItem('shira_msg_templates') || '[]');
  stored.push({ id: 'custom_' + Date.now(), name, text });
  localStorage.setItem('shira_msg_templates', JSON.stringify(stored));
  renderMsgTab();
}

function msgModeChanged() {
  const mode = document.querySelector('input[name="msg-send-mode"]:checked').value;
  document.getElementById('msg-email-row').style.display = (mode === 'b') ? 'block' : 'none';
}

async function sendMessage() {
  const text = document.getElementById('msg-text').value.trim();
  if (!text) { alert('יש להזין תוכן להודעה'); return; }
  if (!selectedCase) { alert('לא נבחר תיק'); return; }

  const mode   = document.querySelector('input[name="msg-send-mode"]:checked').value;
  const btn    = document.getElementById('msg-send-btn');
  const status = document.getElementById('msg-status');
  btn.disabled = true;
  btn.textContent = '⏳ שולח...';
  status.textContent = '';

  const fileId   = selectedCase.fileId || selectedCase.fileMainId;
  const caseData = {
    fileId:     fileId,
    fileNumber: selectedCase.fullFileMainNumber || selectedCase.fileNumber || '',
    sideA:      selectedCase.sideA || '',
    sideB:      selectedCase.sideB || '',
    subject:    selectedCase.subjectSubName || '',
    courtName:  userCourtName || 'בית הדין הרבני',
    courtId:    userCourtId || 5
  };

  try {
    if (mode === 'b') {
      // Option B — direct email, no Shira document
      const toEmail = (document.getElementById('msg-to-email').value || '').trim();
      if (!toEmail || !toEmail.includes('@')) { alert('יש להזין כתובת אימייל תקינה עבור אופן ב'); btn.disabled = false; btn.textContent = '📨 שלח הודעה'; return; }
      const msgSubject = `הודעה מבית הדין הרבני — תיק ${caseData.fileNumber}`;
      const r = await fetch(`${PROXY}/api/send-message-direct`, {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ text, toEmail, subject: msgSubject, caseData })
      });
      const data = await r.json();
      if (data.ok) {
        status.innerHTML = '<span style="color:green">✅ האימייל נשלח ישירות — לא נוצר מסמך בתיק</span>';
      } else {
        status.innerHTML = `<span style="color:red">❌ שגיאה: ${data.error || 'לא ידוע'}</span>`;
      }
    } else {
      // Option A — create document + open Postal
      const r = await fetch(`${PROXY}/api/send-message`, {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ text, caseData })
      });
      const data = await r.json();
      if (data.postalUrl) {
        status.innerHTML = '<span style="color:green">✅ המסמך נוצר בהצלחה — פותח מסך דיוור...</span>';
        window.open(data.postalUrl, '_blank');
      } else {
        status.innerHTML = `<span style="color:red">❌ שגיאה: ${data.error || 'לא ידוע'}</span>`;
      }
    }
  } catch(e) {
    status.innerHTML = `<span style="color:red">❌ שגיאת תקשורת: ${e.message}</span>`;
  } finally {
    btn.disabled = false;
    btn.textContent = '📨 שלח הודעה';
  }
}

// ── Render docs ───────────────────────────────────────────────────────────────
function renderDocs(highlight) {
  const el = document.getElementById('tab-docs');
  if (!caseDocs.length) {
    el.innerHTML = '<p class="empty"><span class="spinner"></span> טוען מסמכים...</p>'; return;
  }
  let html = '';
  caseDocs.forEach(d => {
    const matched = highlight && (d.name.includes(highlight) || (docTexts[d.docId]||'').includes(highlight));
    const matchBadge = matched ? `<span class="match-badge">מכיל "${highlight}"</span>` : '';
    const name = highlight
      ? d.name.replace(new RegExp(highlight,'g'), `<mark>${highlight}</mark>`)
      : d.name;
    html += `<div class="doc-row">
      <span style="font-size:18px">${d.type==='pdf'?'📕':'📄'}</span>
      <span class="doc-name">${name}</span>
      ${matchBadge}
      <span class="doc-type">${d.type.toUpperCase()}</span>
      <span class="doc-date">${d.date}</span>
      <button class="sm primary" onclick="openDoc('${encodeURIComponent(d.openUrl)}')">פתח</button>
    </div>`;
  });
  el.innerHTML = html || '<p class="empty">אין מסמכים</p>';
}

function openDoc(encodedUrl) { window.open(decodeURIComponent(encodedUrl), '_blank'); }

// ── Hearings ──────────────────────────────────────────────────────────────────
async function loadHearings() {
  const el = document.getElementById('tab-hearings');
  el.innerHTML = '<p class="empty"><span class="spinner"></span> טוען דיונים...</p>';
  const fileId = selectedCase.fileId || selectedCase.fileMainId;
  try {
    const r = await fetch(`${PROXY}/api/hearings/${fileId}`);
    const data = await r.json();
    if (data.error || !data.length) { el.innerHTML = '<p class="empty">לא נמצאו דיונים</p>'; return; }
    const today = new Date();
    let html = '';
    data.forEach(row => {
      const dateStr = row.date || row.MeetingDate || '';
      const d    = new Date(dateStr);
      const diff = Math.ceil((d - today) / 86400000);
      const soon = diff >= 0 && diff <= 7 ? `<span class="soon">בעוד ${diff} ימים</span>` : '';
      html += `<div class="hearing-row">
        <span class="hearing-date">${dateStr}</span>
        ${soon}
        <span style="color:#555">${row.type||''}</span>
        <span style="color:#888;font-size:12px">${row.panel||''}</span>
      </div>`;
    });
    el.innerHTML = html;
  } catch { el.innerHTML = `<div class="error-msg">שגיאה בטעינת דיונים</div>`; }
}

// ── Search in case ────────────────────────────────────────────────────────────
function renderSearchTab() {
  document.getElementById('tab-search').innerHTML = `
    <div style="display:flex;gap:4px;margin-bottom:14px;border-bottom:1px solid #e0e4ea;padding-bottom:0;">
      <button class="tab active" id="ssub-case" onclick="switchSearchSubTab('case')">תיק נוכחי בלבד</button>
      <button class="tab"        id="ssub-all"  onclick="switchSearchSubTab('all')">כל תיקי הצדדים</button>
    </div>
    <div id="search-sub-case">
      <div class="row">
        <input type="text" id="content-q" placeholder="חפש מילה או ביטוי בכל מסמכי התיק..." />
        <button onclick="doContentSearch()">🔎 חפש</button>
      </div>
    </div>
    <div id="search-sub-all" style="display:none">
      <div style="background:#f8f9fb;border-radius:8px;padding:10px 14px;margin-bottom:12px;font-size:12px;color:#555;">
        מחפש בכל התיקים של: <strong>${selectedCase.sideA || ''}</strong> ו-<strong>${selectedCase.sideB || ''}</strong>
      </div>
      <div class="row">
        <input type="text" id="content-q-all" placeholder="חפש מילה או ביטוי בכל התיקים..." />
        <button onclick="doContentSearchAll()">🔎 חפש בכולם</button>
      </div>
    </div>
    <div id="content-results"></div>
  `;
}

function switchSearchSubTab(tab) {
  document.getElementById('search-sub-case').style.display = tab === 'case' ? 'block' : 'none';
  document.getElementById('search-sub-all').style.display  = tab === 'all'  ? 'block' : 'none';
  document.getElementById('ssub-case').classList.toggle('active', tab === 'case');
  document.getElementById('ssub-all').classList.toggle('active',  tab === 'all');
  document.getElementById('content-results').innerHTML = '';
}

async function doContentSearch() {
  const q    = document.getElementById('content-q').value.trim();
  const area = document.getElementById('content-results');
  if (!q) return;
  area.innerHTML = '<p class="empty"><span class="spinner"></span> מחפש בתוכן המסמכים...</p>';
  for (const doc of caseDocs) {
    if (!docTexts[doc.docId]) {
      try {
        const r = await fetch(`${PROXY}/api/doctext/${doc.docId}`);
        const d = await r.json();
        docTexts[doc.docId] = d.text || '';
      } catch { docTexts[doc.docId] = ''; }
    }
  }
  const hits = caseDocs.filter(d => d.name.includes(q) || (docTexts[d.docId]||'').includes(q));
  if (!hits.length) { area.innerHTML = `<p class="empty">לא נמצאו תוצאות עבור "${q}"</p>`; return; }
  renderSearchHits(hits, q, area, selectedCase.fileNumber);
}

async function doContentSearchAll() {
  const q    = document.getElementById('content-q-all').value.trim();
  const area = document.getElementById('content-results');
  if (!q) return;

  // Find all cases with same parties from window._cases
  const allCases = (window._cases || []).filter(c =>
    c.sideA === selectedCase.sideA || c.sideB === selectedCase.sideB ||
    c.sideA === selectedCase.sideB || c.sideB === selectedCase.sideA
  );

  if (!allCases.length) {
    area.innerHTML = '<p class="empty">לא נמצאו תיקים נוספים לאותם צדדים. חפש תחילה לפי ת"ז או שם.</p>';
    return;
  }

  area.innerHTML = `<p class="empty"><span class="spinner"></span> טוען מסמכים מ-${allCases.length} תיקים...</p>`;

  let totalHits = 0;
  let html = '';

  for (const c of allCases) {
    const fileId = c.fileId || c.fileMainId;
    // Load docs for this case if not already loaded
    let docs = window._allCaseDocs?.[fileId];
    if (!docs) {
      try {
        const r = await fetch(`${PROXY}/api/documents/${fileId}`);
        docs = await r.json();
        if (!window._allCaseDocs) window._allCaseDocs = {};
        window._allCaseDocs[fileId] = docs;
      } catch { docs = []; }
    }

    // Load doc texts and search
    const hits = [];
    for (const doc of docs) {
      const key = `${fileId}_${doc.docId}`;
      if (!docTexts[key]) {
        try {
          const r = await fetch(`${PROXY}/api/doctext/${doc.docId}`);
          const d = await r.json();
          docTexts[key] = d.text || '';
        } catch { docTexts[key] = ''; }
      }
      if (doc.name.includes(q) || (docTexts[key]||'').includes(q)) {
        hits.push({ ...doc, _textKey: key });
      }
    }

    if (hits.length) {
      totalHits += hits.length;
      html += `<div style="margin-bottom:14px;">
        <div style="font-size:12px;font-weight:600;color:#1a3a5c;padding:6px 0;border-bottom:1px solid #e0e4ea;margin-bottom:8px;">
          📁 תיק ${c.fileNumber} — ${c.subjectSubName || ''}
          <span style="font-weight:400;color:#888;margin-right:8px;">${hits.length} תוצאות</span>
        </div>`;
      hits.forEach(doc => {
        const text    = docTexts[doc._textKey] || '';
        const idx     = text.indexOf(q);
        const snippet = idx >= 0
          ? '...' + text.substring(Math.max(0, idx-40), idx + q.length + 80)
              .replace(new RegExp(q, 'g'), `<mark>${q}</mark>`) + '...'
          : '';
        html += `<div class="doc-row" style="flex-direction:column;align-items:flex-start;gap:4px;">
          <div style="display:flex;gap:8px;align-items:center;width:100%">
            <span>${doc.type==='pdf'?'📕':'📄'}</span>
            <strong style="font-size:13px">${doc.name.replace(new RegExp(q,'g'),`<mark>${q}</mark>`)}</strong>
            <span class="doc-date">${doc.date}</span>
            <button class="sm primary" style="margin-right:auto" onclick="openDoc('${encodeURIComponent(doc.openUrl)}')">פתח</button>
          </div>
          ${snippet ? `<p style="font-size:12px;color:#555;padding-right:24px;line-height:1.7">${snippet}</p>` : ''}
        </div>`;
      });
      html += '</div>';
    }

    // Update progress
    area.innerHTML = `<p class="empty"><span class="spinner"></span> בודק תיק ${c.fileNumber}...</p>`;
  }

  if (!totalHits) {
    area.innerHTML = `<p class="empty">לא נמצאו תוצאות עבור "${q}" בכל ${allCases.length} התיקים</p>`;
    return;
  }
  area.innerHTML = `<p style="font-size:12px;color:#888;margin-bottom:12px;">נמצאו ${totalHits} מסמכים ב-${allCases.length} תיקים</p>` + html;
}

function renderSearchHits(hits, q, area, caseNumber) {
  let html = `<p style="font-size:12px;color:#888;margin-bottom:8px;">נמצאו ${hits.length} מסמכים:</p>`;
  hits.forEach(d => {
    const textKey = d._textKey || d.docId;
    const text    = docTexts[textKey] || docTexts[d.docId] || '';
    const idx     = text.indexOf(q);
    const snippet = idx >= 0
      ? '...' + text.substring(Math.max(0,idx-40), idx+q.length+80)
          .replace(new RegExp(q,'g'), `<mark>${q}</mark>`) + '...'
      : '';
    html += `<div class="doc-row" style="flex-direction:column;align-items:flex-start;gap:4px;">
      <div style="display:flex;gap:8px;align-items:center;width:100%">
        <span>${d.type==='pdf'?'📕':'📄'}</span>
        <strong style="font-size:13px">${d.name}</strong>
        <span class="doc-date">${d.date}</span>
        <button class="sm primary" style="margin-right:auto" onclick="openDoc('${encodeURIComponent(d.openUrl)}')">פתח</button>
      </div>
      ${snippet ? `<p style="font-size:12px;color:#555;padding-right:24px;line-height:1.7">${snippet}</p>` : ''}
    </div>`;
  });
  area.innerHTML = html;
}

// ── AI tab ────────────────────────────────────────────────────────────────────
let aiAbortController = null;
let aiSelectedDocs = new Set(); // manually selected doc IDs

function renderAITab() {
  const el = document.getElementById('tab-ai');
  if (el.innerHTML) return;

  const caseKey = `ai_${selectedCase.fullFileMainNumber || selectedCase.fileNumber}`;

  el.innerHTML = `
    <div style="margin-bottom:12px;">
      <p style="font-size:12px;color:#888;margin-bottom:8px;">שאלה חופשית על התיק</p>

      <!-- Document selector panel -->
      <div style="border:1px solid #e0e4ea;border-radius:8px;padding:12px;margin-bottom:12px;background:#f8f9fb;">
        <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:8px;flex-wrap:wrap;gap:6px;">
          <span style="font-size:12px;font-weight:600;color:#1a3a5c;">📋 בחר מסמכים לסיכום</span>
          <div style="display:flex;gap:6px;flex-wrap:wrap;">
            <button class="sm" onclick="aiSelectRecent(5)"  style="font-size:11px;">5 אחרונים</button>
            <button class="sm" onclick="aiSelectRecent(10)" style="font-size:11px;">10 אחרונים</button>
            <button class="sm" onclick="aiSelectRecent(20)" style="font-size:11px;">20 אחרונים</button>
            <button class="sm" onclick="aiSelectAll()"      style="font-size:11px;background:#e8f5e9;border-color:#c8e6c9;color:#2e7d32;">כל התיק ⚠️</button>
            <button class="sm" onclick="aiSelectNone()"     style="font-size:11px;color:#888;">נקה</button>
          </div>
        </div>
        <input type="text" id="ai-doc-filter" placeholder="סנן לפי שם מסמך..." oninput="renderAIDocList()"
          style="width:100%;margin-bottom:8px;height:32px;font-size:12px;" />
        <div id="ai-doc-list" style="max-height:200px;overflow-y:auto;border:1px solid #e0e4ea;border-radius:6px;background:#fff;"></div>
        <div id="ai-doc-count" style="font-size:11px;color:#1a3a5c;margin-top:6px;font-weight:600;"></div>
      </div>

      <!-- Question input -->
      <div class="row">
        <input type="text" id="ai-q" placeholder='לדוגמה: "מה הסוגיות המרכזיות?" או "סכם את הפרוטוקול"' />
        <button class="primary" id="ai-ask-btn" onclick="askAI()">✨ שאל</button>
        <button class="sm" id="ai-stop-btn" onclick="stopAI()" style="display:none;background:#c62828;color:#fff;border-color:#c62828;">⏹ עצור</button>
      </div>
      <div id="ai-filter-info" style="font-size:11px;color:#1a3a5c;margin-top:6px;display:none;"></div>
    </div>
    <div id="ai-progress-wrap" style="display:none;margin-top:10px;">
      <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:4px;">
        <span id="ai-progress-text" style="font-size:11px;color:#555;"></span>
      </div>
      <div style="background:#e0e4ea;border-radius:4px;height:6px;overflow:hidden;">
        <div id="ai-progress-bar" style="height:6px;background:#1a3a5c;border-radius:4px;width:0%;transition:width 0.3s ease;"></div>
      </div>
    </div>
    <div id="ai-ans-wrap" style="margin-top:14px">
      <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:6px;">
        <p style="font-size:12px;color:#888;">תשובה</p>
        <div style="display:flex;gap:8px;align-items:center;">
          <span class="ai-cost" style="font-size:11px;color:#bbb;display:none;"></span>
          <button class="sm" onclick="exportDocx()" style="font-size:11px;background:#1a3a5c;color:#fff;border-color:#1a3a5c;">📄 ייצא Word</button>
          <button class="sm" onclick="clearAIAnswer('${caseKey}')" style="font-size:11px;color:#aaa;border-color:#e0e4ea;">🗑 נקה</button>
        </div>
      </div>
      <div class="ai-box" id="ai-ans" style="min-height:120px;color:#aaa;font-style:italic;">התשובה תופיע כאן...</div>
    </div>
  `;

  // Init doc list
  aiSelectedDocs = new Set();
  renderAIDocList();
  aiSelectRecent(10); // default: last 10

  // Restore saved answer if exists
  const saved = sessionStorage.getItem(caseKey);
  if (saved) {
    const ans = document.getElementById('ai-ans');
    ans.style.color = '#222';
    ans.style.fontStyle = 'normal';
    ans.textContent = saved;
  }

  // Re-apply dev mode visibility
  if (devMode) {
    const costEl = el.querySelector('.ai-cost');
    if (costEl) costEl.style.display = 'inline';
  }
}

// ── AI doc selector helpers ───────────────────────────────────────────────────
function renderAIDocList() {
  const filter = (document.getElementById('ai-doc-filter')?.value || '').toLowerCase();
  const list   = document.getElementById('ai-doc-list');
  const count  = document.getElementById('ai-doc-count');
  if (!list) return;

  const filtered = caseDocs.filter(d => !filter || d.name.toLowerCase().includes(filter) || d.date.includes(filter));

  if (!filtered.length) {
    list.innerHTML = '<p style="color:#aaa;font-size:12px;padding:8px;">אין מסמכים</p>';
    return;
  }

  list.innerHTML = filtered.map(d => {
    const checked = aiSelectedDocs.has(d.docId) ? 'checked' : '';
    const preloaded = docTexts[d.docId] ? '✓' : '';
    const preloadColor = docTexts[d.docId] ? 'color:#4caf50' : 'color:#ccc';
    return `<label style="display:flex;align-items:center;gap:8px;padding:5px 8px;cursor:pointer;border-bottom:1px solid #f0f2f5;font-size:12px;" onchange="aiToggleDoc('${d.docId}')">
      <input type="checkbox" ${checked} style="accent-color:#1a3a5c;width:14px;height:14px;" />
      <span style="flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;">${d.name}</span>
      <span style="color:#aaa;font-size:11px;white-space:nowrap;">${d.date}</span>
      <span style="font-size:10px;${preloadColor}">${preloaded}</span>
    </label>`;
  }).join('');

  updateAIDocCount();
}

function aiToggleDoc(docId) {
  if (aiSelectedDocs.has(docId)) aiSelectedDocs.delete(docId);
  else aiSelectedDocs.add(docId);
  updateAIDocCount();
}

function updateAIDocCount() {
  const count = document.getElementById('ai-doc-count');
  if (!count) return;
  const n = aiSelectedDocs.size;
  const warn = n > 50 ? ' — ⚠️ סיכום ייקח זמן רב' : n > 20 ? ' — עיבוד בינוני' : '';
  count.textContent = n > 0 ? `נבחרו ${n} מסמכים מתוך ${caseDocs.length}${warn}` : 'לא נבחרו מסמכים';
  count.style.color = n > 50 ? '#e65100' : n > 0 ? '#1a3a5c' : '#aaa';
}

function aiSelectRecent(n) {
  aiSelectedDocs = new Set(caseDocs.slice(0, n).map(d => d.docId));
  renderAIDocList();
}

function aiSelectAll() {
  if (!confirm(`בחירת כל ${caseDocs.length} המסמכים עשויה לקחת זמן רב. להמשיך?`)) return;
  aiSelectedDocs = new Set(caseDocs.map(d => d.docId));
  renderAIDocList();
}

function aiSelectNone() {
  aiSelectedDocs = new Set();
  renderAIDocList();
}

async function loadDocTexts() {
  for (const doc of caseDocs) {
    if (!docTexts[doc.docId]) {
      try {
        const r = await fetch(`${PROXY}/api/doctext/${doc.docId}`);
        const d = await r.json();
        docTexts[doc.docId] = d.text || '';
      } catch { docTexts[doc.docId] = ''; }
    }
  }
}

// ── Extract dates from question ───────────────────────────────────────────────
function extractDatesFromQuestion(q) {
  const dates = [];
  // Match DD/MM/YYYY, D/M/YYYY, DD/MM/YY
  const re = /(\d{1,2})[\/\-\.](\d{1,2})[\/\-\.](\d{2,4})/g;
  let m;
  while ((m = re.exec(q)) !== null) {
    let y = m[3];
    if (y.length === 2) y = '20' + y;
    // Normalize to DD/MM/YYYY
    const day   = m[1].padStart(2, '0');
    const month = m[2].padStart(2, '0');
    dates.push(`${day}/${month}/${y}`);
  }
  return dates;
}

// ── Filter docs by date or keyword ───────────────────────────────────────────
function filterDocsForAI(q) {
  const dates = extractDatesFromQuestion(q);
  const MAX_DOCS = 20; // Gemini supports large context

  // Check for multiple doc types mentioned
  const docTypes = ['פרוטוקול', 'החלטה', 'כתב תביעה', 'כתב הגנה', 'תצהיר', 'חוות דעת', 'הסכם', 'פסק דין'];

  // Filter by date(s) first
  if (dates.length > 0) {
    const filtered = caseDocs.filter(d => dates.some(dt => d.date === dt));
    if (filtered.length > 0) {
      // Within date filter, further filter by doc type if mentioned
      for (const kw of docTypes) {
        if (q.includes(kw)) {
          const typed = filtered.filter(d => d.name.includes(kw));
          if (typed.length > 0) {
            return { docs: typed.slice(0, MAX_DOCS), filterDesc: `סינון: תאריך ${dates.join(', ')} + סוג "${kw}" — ${typed.length} מסמך/ים` };
          }
        }
      }
      return { docs: filtered.slice(0, MAX_DOCS), filterDesc: `סינון לפי תאריך: ${dates.join(', ')} — ${filtered.length} מסמך/ים` };
    }
    return { docs: [], filterDesc: `⚠️ לא נמצאו מסמכים בתאריך ${dates.join(', ')}` };
  }

  // No date — check for doc type keywords
  const matchedTypes = [];
  for (const kw of docTypes) {
    if (q.includes(kw)) {
      const filtered = caseDocs.filter(d => d.name.includes(kw));
      matchedTypes.push(...filtered);
    }
  }
  if (matchedTypes.length > 0) {
    const unique = [...new Map(matchedTypes.map(d => [d.docId, d])).values()];
    return { docs: unique.slice(0, MAX_DOCS), filterDesc: `סינון לפי סוג מסמך — ${unique.length} מסמך/ים` };
  }

  // Check for "כל התיק" or "סכם הכל"
  if (q.includes('כל התיק') || q.includes('סכם הכל') || q.includes('כל המסמכים')) {
    return { docs: caseDocs.slice(0, MAX_DOCS), filterDesc: `כל התיק — ${Math.min(caseDocs.length, MAX_DOCS)} מסמכים (מתוך ${caseDocs.length})` };
  }

  // Default — most recent docs
  const recent = [...caseDocs].slice(0, 10);
  return {
    docs: recent,
    filterDesc: caseDocs.length > 10
      ? `⚠️ ${caseDocs.length} מסמכים בתיק — טוען 10 האחרונים. ציין תאריך, סוג, או "כל התיק" לשליטה מדויקת.`
      : `${caseDocs.length} מסמכים בתיק`
  };
}

// ── Stop AI ───────────────────────────────────────────────────────────────────
function stopAI() {
  if (aiAbortController) {
    aiAbortController.abort();
    aiAbortController = null;
  }
  const ans = document.getElementById('ai-ans');
  if (ans) {
    const current = ans.textContent.replace('', '');
    ans.textContent = current ? current + '\\n\\n[הופסק על ידי המשתמש]' : '[הופסק]';
  }
  document.getElementById('ai-ask-btn').style.display = '';
  document.getElementById('ai-stop-btn').style.display = 'none';
}

async function askAI() {
  const q = document.getElementById('ai-q').value.trim();
  if (!q) return;

  const ans = document.getElementById('ai-ans');
  const askBtn  = document.getElementById('ai-ask-btn');
  const stopBtn = document.getElementById('ai-stop-btn');
  const filterInfo = document.getElementById('ai-filter-info');
  const progressWrap = document.getElementById('ai-progress-wrap');
  const progressBar  = document.getElementById('ai-progress-bar');
  const progressText = document.getElementById('ai-progress-text');

  // Use manually selected docs, fall back to smart filter if none selected
  let docsToUse, filterDesc;
  if (aiSelectedDocs && aiSelectedDocs.size > 0) {
    docsToUse  = caseDocs.filter(d => aiSelectedDocs.has(d.docId));
    filterDesc = `נבחרו ידנית ${docsToUse.length} מסמכים`;
  } else {
    const result = filterDocsForAI(q);
    docsToUse  = result.docs;
    filterDesc = result.filterDesc;
  }

  filterInfo.textContent = filterDesc;
  filterInfo.style.display = 'block';

  if (docsToUse.length === 0) {
    ans.className = 'ai-box';
    ans.style.color = '#c62828';
    ans.textContent = 'לא נבחרו מסמכים — בחר מסמכים מהרשימה למעלה';
    return;
  }

  askBtn.style.display = 'none';
  stopBtn.style.display = '';
  stopBtn.disabled = false;

  progressWrap.style.display = 'block';
  progressBar.style.width = '0%';
  ans.className = 'ai-box loading';
  ans.textContent = '';

  const toLoad = docsToUse.filter(d => !docTexts[d.docId]);
  const totalToLoad = toLoad.length;
  let loaded = 0;

  for (const doc of docsToUse) {
    if (!docTexts[doc.docId]) {
      try {
        loaded++;
        const pct = Math.round((loaded / Math.max(totalToLoad, 1)) * 80);
        progressBar.style.width = pct + '%';
        progressText.textContent = `טוען מסמך ${loaded} מתוך ${totalToLoad}: ${doc.name}`;
        const controller = new AbortController();
        const timeoutId = setTimeout(() => controller.abort(), 8000);
        try {
          const r = await fetch(`${PROXY}/api/doctext/${doc.docId}`, { signal: controller.signal });
          clearTimeout(timeoutId);
          const d = await r.json();
          docTexts[doc.docId] = d.text || '';
        } catch {
          clearTimeout(timeoutId);
          docTexts[doc.docId] = '';
        }
      } catch { docTexts[doc.docId] = ''; }
    }
  }

  progressBar.style.width = '95%';
  progressText.textContent = '';

  const combined = docsToUse
    .map(d => `[${d.name} | ${d.date}]:\\n${(docTexts[d.docId]||'')}`)
    .join('\\n\\n');

  try {
    await streamAI(
      'אתה עוזר משפטי לבית הדין הרבני. ענה בעברית בלבד. ענה על בסיס המסמכים שסופקו לך בלבד. חשוב מאוד: שם הקובץ אינו רלוונטי — גם אם המסמך נקרא "כריכה" או כל שם אחר, הוא עשוי להכיל פרוטוקול דיון, החלטה, עדות או כל תוכן משפטי אחר. סכם את התוכן בפועל של המסמך ללא קשר לשמו. כתוב טקסט רציף ונקי ללא סימני markdown, ללא כוכביות, ללא hashtag, ללא מקפים כסמני רשימה. השתמש במספור רגיל (1. 2. 3.) לרשימות. הפרד בין פסקאות בשורה ריקה בלבד.',
      `מסמכי תיק:\\n\\n${combined}\\n\\n---\\nשאלה: ${q}`,
      ans
    );
  } catch(e) {
    if (e.name === 'AbortError') {
      // Already handled in stopAI
    } else {
      ans.className = 'ai-box';
      ans.textContent = 'שגיאה: ' + e.message;
    }
  } finally {
    askBtn.style.display = '';
    stopBtn.style.display = 'none';
    aiAbortController = null;
    // Hide progress bar
    const pw = document.getElementById('ai-progress-wrap');
    const pb = document.getElementById('ai-progress-bar');
    if (pw) {
      pb.style.width = '100%';
      setTimeout(() => { pw.style.display = 'none'; pb.style.width = '0%'; }, 600);
    }
  }
}
// ── Streaming AI call ─────────────────────────────────────────────────────────
async function streamAI(system, userMessage, targetEl) {
  targetEl.className = 'ai-box';
  targetEl.innerHTML = '<span class="ai-cursor"></span>';

  aiAbortController = new AbortController();

  const resp = await fetch(`${PROXY}/api/ai`, {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ system, messages: [{ role: 'user', content: userMessage }] }),
    signal: aiAbortController.signal
  });

  if (!resp.ok) throw new Error(`HTTP ${resp.status}`);

  const reader  = resp.body.getReader();
  const decoder = new TextDecoder();
  let text = '';

  while (true) {
    const { done, value } = await reader.read();
    if (done) break;
    const lines = decoder.decode(value, { stream: true }).split('\\n');
    for (const line of lines) {
      if (!line.startsWith('data:')) continue;
      const raw = line.slice(5).trim();
      if (!raw || raw === '[DONE]') continue;
      try {
        const chunk = JSON.parse(raw);
        if (chunk.error) throw new Error(chunk.error);
        if (chunk.text) {
          text += chunk.text;
          targetEl.innerHTML = text + '<span class="ai-cursor"></span>';
        }
        if (chunk.usage) {
          const u = chunk.usage;
          const costEl = targetEl.parentElement.querySelector('.ai-cost');
          if (costEl) {
            costEl.textContent = `טוקנים: ${u.total.toLocaleString()} | עלות: $${u.usd} (~₪${u.ils})`;
          }
        }
        if (chunk.done) break;
      } catch(e) { /* skip malformed chunks */ }
    }
  }

  // Final render without cursor
  targetEl.style.color = '#222';
  targetEl.style.fontStyle = 'normal';
  targetEl.textContent = text || 'לא ניתן לענות';

  // Save to sessionStorage keyed by case number
  if (selectedCase) {
    const caseKey = `ai_${selectedCase.fullFileMainNumber || selectedCase.fileNumber}`;
    try { sessionStorage.setItem(caseKey, text); } catch(e) {}
  }

  return text;
}

function clearAIAnswer(caseKey) {
  sessionStorage.removeItem(caseKey);
  const ans = document.getElementById('ai-ans');
  if (ans) {
    ans.style.color = '#aaa';
    ans.style.fontStyle = 'italic';
    ans.textContent = 'התשובה תופיע כאן...';
  }
  const cost = document.querySelector('.ai-cost');
  if (cost) cost.textContent = '';
}

async function exportDocx() {
  const ans = document.getElementById('ai-ans');
  if (!ans || !ans.textContent || ans.textContent === 'התשובה תופיע כאן...') {
    alert('אין תשובה לייצוא — הפעל שאילתת AI תחילה.');
    return;
  }
  try {
    const resp = await fetch(`${PROXY}/api/export-docx`, {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({
        text:        ans.textContent,
        caseNumber:  selectedCase?.fullFileMainNumber || selectedCase?.fileNumber || '',
        caseTitle:   `${selectedCase?.sideA || ''} נגד ${selectedCase?.sideB || ''}`,
        courtName:   userCourtName || 'בית הדין הרבני'
      })
    });
    if (!resp.ok) { alert('שגיאה בייצוא המסמך'); return; }
    const blob = await resp.blob();
    const url  = URL.createObjectURL(blob);
    const a    = document.createElement('a');
    const cd   = resp.headers.get('Content-Disposition') || '';
    const match = cd.match(/filename\\*?=(?:UTF-8'')?([^;]+)/i);
    a.download = match ? decodeURIComponent(match[1]) : 'סיכום_AI.docx';
    a.href = url;
    a.click();
    URL.revokeObjectURL(url);
  } catch(e) {
    alert('שגיאה: ' + e.message);
  }
}

async function askAI() {
  const q = document.getElementById('ai-q').value.trim();
  if (!q) return;

  const ans = document.getElementById('ai-ans');
  ans.className = 'ai-box loading';
  ans.textContent = 'טוען מסמכים...';

  // Load doc texts if not already loaded
  await loadDocTexts();

  const combined = caseDocs
    .map(d => `[${d.name}]:\\n${(docTexts[d.docId]||'').substring(0,30000)}`)
    .join('\\n\\n');
  const context = combined || `תיק ${selectedCase.subjectSubName}`;
  try {
    await streamAI(
      'אתה עוזר משפטי לבית הדין הרבני. ענה בעברית בלבד. ענה על בסיס המסמכים בלבד. כתוב טקסט רציף ונקי ללא סימני markdown, ללא כוכביות, ללא hashtag, ללא מקפים כסמני רשימה. השתמש במספור רגיל (1. 2. 3.) לרשימות. הפרד בין פסקאות בשורה ריקה בלבד.',
      `מסמכי תיק:\\n\\n${context}\\n\\n---\\nשאלה: ${q}`,
      ans
    );
  } catch(e) {
    ans.className = 'ai-box';
    ans.textContent = 'שגיאה: ' + e.message;
  }
}

// ── Keyboard shortcuts ────────────────────────────────────────────────────────
document.addEventListener('keydown', e => {
  if (e.key === 'Enter') {
    const id = document.activeElement.id;
    if (id === 'id-input')   doSearch();
    if (id === 'case-input') doCaseSearch();
    if (id === 'name-last' || id === 'name-first') doNameSearch();
    if (id === 'content-q')     doContentSearch();
    if (id === 'content-q-all') doContentSearchAll();
    if (id === 'search-all-q')  doSearchAllCases();
    if (id === 'ai-q')          askAI();
  }
});

// ── Usage stats (dev mode only) ───────────────────────────────────────────────
async function showUsage() {
  const popup   = document.getElementById('usage-popup');
  const content = document.getElementById('usage-content');
  popup.style.display = popup.style.display === 'none' ? 'block' : 'none';
  if (popup.style.display === 'none') return;
  content.textContent = 'טוען...';
  try {
    const r = await fetch(`${PROXY}/api/usage`);
    const d = await r.json();
    if (d.error) { content.textContent = 'שגיאה: ' + d.error; return; }
    content.innerHTML = `
      <div style="display:grid;grid-template-columns:1fr 1fr;gap:8px;margin-bottom:12px;">
        <div style="background:#f8f9fb;border-radius:8px;padding:10px;text-align:center;">
          <div style="font-size:20px;font-weight:600;color:#1a3a5c;">${d.queries}</div>
          <div style="font-size:11px;color:#888;">שאילתות</div>
        </div>
        <div style="background:#f8f9fb;border-radius:8px;padding:10px;text-align:center;">
          <div style="font-size:20px;font-weight:600;color:#1a3a5c;">${d.total_tokens.toLocaleString()}</div>
          <div style="font-size:11px;color:#888;">סה"כ טוקנים</div>
        </div>
        <div style="background:#e8f5e9;border-radius:8px;padding:10px;text-align:center;">
          <div style="font-size:20px;font-weight:600;color:#2e7d32;">$${d.total_usd}</div>
          <div style="font-size:11px;color:#888;">סה"כ עלות</div>
        </div>
        <div style="background:#e8f5e9;border-radius:8px;padding:10px;text-align:center;">
          <div style="font-size:20px;font-weight:600;color:#2e7d32;">₪${d.total_ils}</div>
          <div style="font-size:11px;color:#888;">בשקלים</div>
        </div>
      </div>
      <div style="font-size:11px;color:#888;">ממוצע לשאילתה: $${d.avg_usd}</div>
      ${d.last_queries.length ? `
        <div style="margin-top:12px;font-size:11px;color:#888;border-top:1px solid #f0f2f5;padding-top:10px;">10 שאילתות אחרונות:</div>
        ${d.last_queries.slice().reverse().map(q =>
          `<div style="display:flex;justify-content:space-between;font-size:11px;padding:3px 0;border-bottom:1px solid #f8f9fb;">
            <span style="color:#666;">${q.ts.replace('T',' ')}</span>
            <span style="color:#888;">${q.total.toLocaleString()} טוקנים</span>
            <span style="color:#2e7d32;">$${q.usd} / ₪${q.ils}</span>
          </div>`
        ).join('')}
      ` : ''}
    `;
  } catch(e) {
    content.textContent = 'שגיאה בטעינת נתונים';
  }
}

// ── Start ─────────────────────────────────────────────────────────────────────
boot();
</script>
</body>
</html>
"""




SHIRA = "http://shira2"
SPFE  = "http://prod-spfe:1000"
PROXY = "http://192.168.174.80:8080"

# ↓↓↓ PUT YOUR GEMINI API KEY HERE — ONLY HERE ↓↓↓
GEMINI_API_KEY = "AIzaSyCgutrB9sRoyQHC5mY11LiHWF505VQVD44"
# ↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑↑

def make_session():
    s = requests.Session()
    s.auth = HttpNegotiateAuth()
    s.headers.update({
        "Content-Type": "application/json; charset=UTF-8",
        "Origin": SHIRA,
        "Referer": f"{SHIRA}/App/main/files/files-list"
    })
    s.proxies = {"http": None, "https": None}
    return s

SESSION = make_session()


@app.route("/")
def index():
    return Response(_HTML, mimetype="text/html; charset=utf-8")

@app.route("/api/me")
def me():
    court_names = {
        1: "ירושלים", 2: "תל אביב", 3: "חיפה", 4: "פתח תקוה",
        5: "רחובות",  6: "באר שבע", 7: "טבריה", 8: "צפת",
        9: "אשדוד",  10: "אשקלון", 11: "נתניה",
        12: "בית הדין הגדול", 13: "אריאל"
    }
    try:
        r = SESSION.get(
            f"{SHIRA}/api/api/userController/GetUser",
            timeout=10
        )
        r.raise_for_status()
        data = r.json()
        user_id    = data.get("userId")
        user_name  = data.get("userName")
        first_name = data.get("firstName", "")
        last_name  = data.get("lastName", "")
        court_list = data.get("courtList", [])
        if court_list:
            court_id   = court_list[0]["courtId"]
            court_name = court_list[0].get("courtName") or court_names.get(court_id, str(court_id))
            print(f"[me] user={user_name} court={court_id} ({court_name})")
            return jsonify({
                "courtId":   court_id,
                "courtName": court_name,
                "userId":    user_id,
                "userName":  user_name,
                "firstName": first_name,
                "lastName":  last_name
            })
        return jsonify({"error": "no court found for this user"}), 500
    except Exception as e:
        print(f"[me] exception: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/search", methods=["POST"])
def search():
    body = request.json
    id_num = body.get("idNum", "").strip()
    if not id_num:
        return jsonify({"error": "idNum required"}), 400
    payload = {
        "courtID": None, "assemblyId": None, "fileNumber": None,
        "fileMainID": None, "subjectID": None, "subjectSubID": None,
        "Composition": None, "FileStatusOpen": "-1",
        "FirstName": None, "IdNum1": id_num, "IdType1": 1,
        "IsOnlineFile": False, "LastName": None, "OldFileNum": "",
        "currentPage": 1, "fileStatusID": None,
        "insertDateFrom": None, "insertDateTo": None,
        "isCorrectName": False, "isPriority": False,
        "meetingDateFrom": None, "meetingDateTo": None,
        "rowsPerPage": 100
    }
    try:
        resp = SESSION.post(f"{SHIRA}/api/api/FileSearch/GetAdvancedFileSearch", json=payload, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        for f in data:
            f["sideB"] = f.get("sideB") or ""
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/search-case", methods=["POST"])
def search_case():
    body = request.json
    file_main_id = body.get("fileMainId", "").strip()
    file_number  = body.get("fileNumber")
    if not file_main_id:
        return jsonify({"error": "fileMainId required"}), 400
    payload = {
        "courtID": None, "assemblyId": None,
        "fileNumber": int(file_number) if file_number else None,
        "fileMainID": int(file_main_id),
        "subjectID": None, "subjectSubID": None,
        "Composition": None, "FileStatusOpen": "-1",
        "FirstName": None, "IdNum1": None, "IdType1": 1,
        "IsOnlineFile": False, "LastName": None, "OldFileNum": "",
        "currentPage": 1, "fileStatusID": None,
        "insertDateFrom": None, "insertDateTo": None,
        "isCorrectName": False, "isPriority": False,
        "meetingDateFrom": None, "meetingDateTo": None,
        "rowsPerPage": 100
    }
    try:
        resp = SESSION.post(f"{SHIRA}/api/api/FileSearch/GetAdvancedFileSearch", json=payload, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        for f in data:
            f["sideB"] = f.get("sideB") or ""
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/search-name", methods=["POST"])
def search_name():
    body       = request.json
    last_name  = body.get("lastName", "").strip()
    first_name = body.get("firstName", "").strip()
    if not last_name and not first_name:
        return jsonify({"error": "lastName or firstName required"}), 400
    payload = {
        "courtID": None, "assemblyId": None, "fileNumber": None,
        "fileMainID": None, "subjectID": None, "subjectSubID": None,
        "Composition": None, "FileStatusOpen": "-1",
        "FirstName": first_name or None,
        "LastName":  last_name  or None,
        "IdNum1": None, "IdType1": 1,
        "IsOnlineFile": False, "OldFileNum": "",
        "currentPage": 1, "fileStatusID": None,
        "insertDateFrom": None, "insertDateTo": None,
        "isCorrectName": False, "isPriority": False,
        "meetingDateFrom": None, "meetingDateTo": None,
        "rowsPerPage": 100
    }
    try:
        resp = SESSION.post(f"{SHIRA}/api/api/FileSearch/GetAdvancedFileSearch", json=payload, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        for f in data:
            f["sideB"] = f.get("sideB") or ""
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 500



def documents(file_id):
    url = f"{SHIRA}/classic/Forms/File/Contents/FileDocs.aspx?userid=0&courtid=0&FileID={file_id}&EntityId={file_id}&EntityTypeId=6"
    try:
        resp = SESSION.get(url, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")
        docs = []
        table = soup.find("table", id="grdFileDocs")
        if not table:
            for t in soup.find_all("table"):
                if "OpenDocument" in str(t):
                    table = t
                    break
        if table:
            for tr in table.find_all("tr"):
                row_html = str(tr)
                doc_id_match = re.search(r"OpenDocument\((\d+)\)", row_html)
                if not doc_id_match:
                    continue
                doc_id = doc_id_match.group(1)
                link = tr.find("a", href=re.compile(r"OpenDocument")) or tr.find("a", onclick=True)
                name = link.get_text(strip=True) if link else f"מסמך {doc_id}"
                if not name:
                    name = f"מסמך {doc_id}"
                row_text = tr.get_text(" ", strip=True)
                date_match = re.search(r"\d{2}/\d{2}/\d{4}", row_text)
                date = date_match.group(0) if date_match else ""
                ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
                docs.append({
                    "docId": doc_id, "name": name, "date": date,
                    "type": "pdf" if ext == "pdf" else "docx",
                    "openUrl": f"{SHIRA}/classic/Forms/Documents/DM/DMOpenDocument.aspx?DocIDs={doc_id}&Action=1"
                })
        return jsonify(docs)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/documents/<int:file_id>", methods=["GET"])
def documents(file_id):
    url = f"{SHIRA}/classic/Forms/File/Contents/FileDocs.aspx?userid=0&courtid=0&FileID={file_id}&EntityId={file_id}&EntityTypeId=6"
    try:
        resp = SESSION.get(url, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")
        docs = []
        table = soup.find("table", id="grdFileDocs")
        if not table:
            for t in soup.find_all("table"):
                if "OpenDocument" in str(t):
                    table = t
                    break
        if table:
            for tr in table.find_all("tr"):
                row_html = str(tr)
                doc_id_match = re.search(r"OpenDocument\((\d+)\)", row_html)
                if not doc_id_match:
                    continue
                doc_id = doc_id_match.group(1)
                link = tr.find("a", href=re.compile(r"OpenDocument")) or tr.find("a", onclick=True)
                name = link.get_text(strip=True) if link else f"מסמך {doc_id}"
                if not name:
                    name = f"מסמך {doc_id}"
                row_text = tr.get_text(" ", strip=True)
                date_match = re.search(r"\d{2}/\d{2}/\d{4}", row_text)
                date = date_match.group(0) if date_match else ""
                ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
                docs.append({
                    "docId": doc_id, "name": name, "date": date,
                    "type": "pdf" if ext == "pdf" else "docx",
                    "openUrl": f"{SHIRA}/classic/Forms/Documents/DM/DMOpenDocument.aspx?DocIDs={doc_id}&Action=1"
                })
        return jsonify(docs)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/doctext/<doc_id>", methods=["GET"])
def doc_text(doc_id):
    try:
        xml = f"<XmlData><DocumentID>{doc_id}</DocumentID></XmlData>"
        r1 = SESSION.post(f"{SHIRA}/classic/WS/App/WsShiraUtils.asmx/GetDocumentDetails",
                          data=xml.encode("utf-8"), headers={"Content-Type": "application/xml"}, timeout=10)
        root1 = ET.fromstring(r1.text)
        doc_number_el = root1.find("DocNumber")
        if doc_number_el is None or not doc_number_el.text:
            return jsonify({"text": "", "error": "DocNumber not found"})
        doc_number = doc_number_el.text.strip()
        r2 = SESSION.post(f"{SPFE}/ShiraDocsMngWS.asmx/GetDocumentUrlAndStatus",
                          data=f"{{'docNumber':'{doc_number}', 'isCopy':'true'}}",
                          headers={"Content-Type": "application/json"}, timeout=10)
        result = r2.json().get("d", "")
        file_url = result.split("|")[0] if "|" in result else result
        if not file_url or file_url == "-1":
            return jsonify({"text": "", "error": "file URL not found"})
        r3 = SESSION.get(file_url, timeout=20)
        r3.raise_for_status()
        file_bytes = io.BytesIO(r3.content)
        ext = file_url.rsplit(".", 1)[-1].lower()
        text = ""
        if ext == "pdf":
            with pdfplumber.open(file_bytes) as pdf:
                text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        elif ext in ("docx", "doc"):
            d = docx.Document(file_bytes)
            text = "\n".join(p.text for p in d.paragraphs)
        else:
            text = r3.content.decode("utf-8", errors="ignore")[:50000]
        return jsonify({"text": text[:30000], "url": file_url})
    except Exception as e:
        return jsonify({"text": "", "error": str(e)})


@app.route("/api/hearings/<int:file_id>", methods=["GET"])
def hearings(file_id):
    url = f"{SHIRA}/classic/Forms/File/Contents/FileMeetings.aspx?userid=0&courtid=0&FileID={file_id}&EntityId={file_id}&EntityTypeId=6"
    try:
        resp = SESSION.get(url, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")
        table = soup.find("table", id="grdMeetings") or soup.find("table")
        if not table:
            return jsonify([])
        rows = []
        for tr in table.find_all("tr")[1:]:
            cells = [td.get_text(strip=True) for td in tr.find_all("td")]
            if len(cells) >= 4 and cells[2]:
                rows.append({
                    "meetingId": cells[0], "hebrewDate": cells[1],
                    "date": cells[2], "type": cells[3],
                    "timeFrom": cells[4] if len(cells) > 4 else "",
                    "timeTo": cells[5] if len(cells) > 5 else "",
                    "panel": cells[6] if len(cells) > 6 else "",
                    "status": cells[7] if len(cells) > 7 else "",
                })
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def anonymize_text(text):
    """
    Remove PII before sending to external AI.
    Removes: ID numbers, phones, emails, bank accounts, addresses, names.
    """
    import re

    # Israeli ID numbers (9 digits)
    text = re.sub(r'\b\d{9}\b', '[ת"ז]', text)

    # Phone numbers
    text = re.sub(r'\b0\d{1,2}[-\s]?\d{3}[-\s]?\d{4}\b', '[טלפון]', text)

    # Email addresses
    text = re.sub(r'[\w.+\-]+@[\w\-]+\.\w+', '[מייל]', text)

    # Bank accounts near keywords
    text = re.sub(r'(?:חשבון|ח-ן)[^\d]{0,10}(\d{5,12})',
                  lambda m: m.group(0).replace(m.group(1), '[חשבון]'), text)

    # Addresses
    text = re.sub(r'(?:רח\'|רחוב|שד\'|שדרות|סמטת|ככר)\s+[\u05d0-\u05ea\s"\']{2,30}\s+\d{1,4}',
                  '[כתובת]', text)

    # Zip codes (7 digits)
    text = re.sub(r'\b\d{7}\b', '[מיקוד]', text)

    # Passport numbers
    text = re.sub(r'\b[A-Z]{1,2}\d{6,8}\b', '[דרכון]', text)

    # Names after "האיש" / "האישה"
    text = re.sub(
        r'(?:האיש|האישה)\s+([\u05d0-\u05ea]{2,10}\s+[\u05d0-\u05ea]{2,10})',
        lambda m: m.group(0).replace(m.group(1), '[שם]'), text)

    # Names after legal titles (attorney, judge, dayan)
    text = re.sub(
        r'(?:עו"ד|עורך דין|עורכת דין|השופט|השופטת|הדיין|הרב)\s+'
        r'([\u05d0-\u05ea]{2,10}(?:\s+[\u05d0-\u05ea]{2,10}){1,2})',
        lambda m: m.group(0).replace(m.group(1), '[שם]'), text)

    # Names after party roles: תובע, נתבע etc.
    text = re.sub(
        r'(?:התובע|התובעת|הנתבע|הנתבעת|המבקש|המבקשת|המשיב|המשיבה),?\s+'
        r'([\u05d0-\u05ea]{2,10}\s+[\u05d0-\u05ea]{2,10})',
        lambda m: m.group(0).replace(m.group(1), '[שם]'), text)

    return text

@app.route("/api/ai-test")
def ai_test():
    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={GEMINI_API_KEY}"
        print(f"[ai-test] calling Gemini...")
        resp = requests.post(
            url,
            json={"contents": [{"parts": [{"text": "say hello in Hebrew"}]}]},
            proxies={"https": None, "http": None},
            verify=False,
            timeout=60
        )
        print(f"[ai-test] status={resp.status_code}")
        print(f"[ai-test] raw={resp.text[:300]}")
        return jsonify({"status": resp.status_code, "raw": resp.text[:500]})
    except Exception as e:
        print(f"[ai-test] exception: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai", methods=["POST"])
def ai_proxy():
    """
    Streaming endpoint — returns Server-Sent Events.
    The frontend should read chunks as they arrive instead of waiting for the full response.
    """
    body = request.json
    user_message = body.get("messages", [{}])[0].get("content", "")
    system = body.get("system", "אתה עוזר משפטי לבית הדין הרבני. ענה בעברית בלבד. כתוב טקסט רציף ונקי ללא סימני markdown, ללא כוכביות, ללא hashtag, ללא מקפים כסמני רשימה. השתמש במספור רגיל (1. 2. 3.) לרשימות. הפרד בין פסקאות בשורה ריקה בלבד.")

    # Safety cap — Gemini 2.5 Flash supports ~1M tokens but keep requests reasonable
    user_message = user_message[:200000]

    print(f"[ai] raw message length BEFORE anonymize={len(user_message)}")
    print(f"[ai] first 500 chars: {repr(user_message[:500])}")

    # Anonymize PII before sending to external AI
    user_message = anonymize_text(user_message)

    # Force Hebrew response — appended after anonymization so it's not stripped
    user_message = user_message + "\n\n[הוראה חובה: ענה תמיד בעברית בלבד, ללא יוצא מן הכלל]"
    print(f"[ai] message anonymized and ready, length={len(user_message)}")

    stream_url = (
        f"https://generativelanguage.googleapis.com/v1beta/models/"
        f"gemini-2.5-flash:streamGenerateContent?alt=sse&key={GEMINI_API_KEY}"
    )

    print(f"[ai] streaming to Gemini, message length={len(user_message)}")

    @stream_with_context
    def generate():
        try:
            with requests.post(
                stream_url,
                json={
                    "contents": [{"parts": [{"text": user_message}]}],
                    "systemInstruction": {"parts": [{"text": system}]}
                },
                proxies={"https": None, "http": None},
                verify=False,
                timeout=180,   # 3 minutes — enough for very large docs
                stream=True
            ) as resp:
                print(f"[ai] gemini stream status={resp.status_code}")
                for raw_line in resp.iter_lines():
                    if not raw_line:
                        continue
                    # SSE lines look like: b"data: {...json...}"
                    if isinstance(raw_line, bytes):
                        raw_line = raw_line.decode("utf-8")
                    if not raw_line.startswith("data:"):
                        continue
                    chunk_str = raw_line[5:].strip()
                    if chunk_str == "[DONE]":
                        break
                    try:
                        chunk_data = json.loads(chunk_str)
                        # Extract text from the chunk
                        parts = (
                            chunk_data.get("candidates", [{}])[0]
                            .get("content", {})
                            .get("parts", [])
                        )
                        for part in parts:
                            text_piece = part.get("text", "")
                            if text_piece:
                                # Strip markdown symbols
                                text_piece = text_piece.replace("**", "").replace("__", "")
                                text_piece = text_piece.replace("*", "").replace("_", "")
                                text_piece = text_piece.replace("##", "").replace("###", "").replace("#", "")
                                # Send as SSE event to the frontend
                                yield f"data: {json.dumps({'text': text_piece}, ensure_ascii=False)}\n\n"

                        # Log token usage when it appears (last chunk)
                        usage = chunk_data.get("usageMetadata")
                        if usage:
                            input_tokens    = usage.get("promptTokenCount", 0)
                            output_tokens   = usage.get("candidatesTokenCount", 0)
                            thinking_tokens = usage.get("thoughtsTokenCount", 0)
                            total_tokens    = usage.get("totalTokenCount", 0)

                            # Gemini 2.5 Flash pricing (USD per 1M tokens)
                            PRICE_INPUT    = 0.15
                            PRICE_OUTPUT   = 0.60
                            PRICE_THINKING = 3.50

                            cost_usd = (
                                (input_tokens    / 1_000_000) * PRICE_INPUT +
                                (output_tokens   / 1_000_000) * PRICE_OUTPUT +
                                (thinking_tokens / 1_000_000) * PRICE_THINKING
                            )
                            cost_ils = cost_usd * 3.7  # approximate ILS rate

                            print(
                                f"[ai] tokens — input: {input_tokens:,}, "
                                f"output: {output_tokens:,}, "
                                f"thinking: {thinking_tokens:,}, "
                                f"total: {total_tokens:,} | "
                                f"cost: ${cost_usd:.4f} (~₪{cost_ils:.3f})"
                            )

                            # Append to usage log file
                            import datetime
                            log_entry = {
                                "ts":       datetime.datetime.now().isoformat(timespec="seconds"),
                                "input":    input_tokens,
                                "output":   output_tokens,
                                "thinking": thinking_tokens,
                                "total":    total_tokens,
                                "usd":      round(cost_usd, 6),
                                "ils":      round(cost_ils, 4),
                            }
                            try:
                                log_path = os.path.join(os.path.dirname(__file__), "usage_log.jsonl")
                                with open(log_path, "a", encoding="utf-8") as lf:
                                    lf.write(json.dumps(log_entry, ensure_ascii=False) + "\n")
                            except Exception as le:
                                print(f"[ai] log write error: {le}")

                            # Send usage summary to frontend
                            yield f"data: {json.dumps({'usage': log_entry})}\n\n"
                    except json.JSONDecodeError:
                        continue

            # Signal end of stream
            yield f"data: {json.dumps({'done': True})}\n\n"

        except Exception as e:
            print(f"[ai] stream exception: {e}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",   # important for nginx proxies
            "Connection": "keep-alive",
        }
    )


@app.route("/api/export-docx", methods=["POST"])
def export_docx():
    """
    Generate a Word document from AI summary text.
    Matches the formatting of the reference document (החלטה_סיכומים.docx).
    """
    import io
    import datetime
    import docx as _docx
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    body = request.json
    text        = body.get("text", "").strip()
    case_number = body.get("caseNumber", "")
    case_title  = body.get("caseTitle", "")
    court_name  = body.get("courtName", "בית הדין הרבני")

    if not text:
        return jsonify({"error": "no text provided"}), 400

    doc = _docx.Document()

    # ── Page setup (match reference: ~A4 with custom margins) ──────────────────
    section = doc.sections[0]
    section.page_width  = 7560310   # ~534 mm  (same as reference)
    section.page_height = 10692130  # ~756 mm
    section.left_margin   = 900430   # ~1.58 cm
    section.right_margin  = 1141095  # ~2.01 cm
    section.top_margin    = 331470   # ~0.58 cm
    section.bottom_margin = 810260   # ~1.43 cm

    # ── RTL document direction ──────────────────────────────────────────────────
    body_el = doc.element.body
    sectPr = body_el.get_or_add_sectPr()
    bidi = OxmlElement('w:bidi')
    sectPr.append(bidi)

    FONT_NAME = 'FrankRuehl'   # Frank Ruehl — standard rabbinical court font
    FONT_SIZE = 14

    def set_rtl_paragraph(p, space_before_pt=4, space_after_pt=4, center=False):
        pf = p.paragraph_format
        pf.space_before = Pt(space_before_pt)
        pf.space_after  = Pt(space_after_pt)
        pf.alignment    = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
        pPr = p._p.get_or_add_pPr()
        bidi_el = OxmlElement('w:bidi')
        pPr.append(bidi_el)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center' if center else 'both')
        pPr.append(jc)
        return p

    def add_paragraph(text_content, bold=False, center=False,
                      space_before_pt=4, space_after_pt=4,
                      font_size=None, color=None, underline=False):
        p = doc.add_paragraph()
        set_rtl_paragraph(p, space_before_pt, space_after_pt, center)
        run = p.add_run(text_content)
        run.font.name      = FONT_NAME
        run.font.size      = Pt(font_size or FONT_SIZE)
        run.font.bold      = bold
        run.font.underline = underline
        if color:
            run.font.color.rgb = RGBColor(*color)
        rPr = run._r.get_or_add_rPr()
        rtl_el = OxmlElement('w:rtl')
        rPr.append(rtl_el)
        lang = OxmlElement('w:lang')
        lang.set(qn('w:bidi'), 'he-IL')
        rPr.append(lang)
        return p

    # ── Header ──────────────────────────────────────────────────────────────────
    add_paragraph("בבית הדין הרבני האזורי", bold=True, center=True, space_before_pt=6, space_after_pt=2)
    add_paragraph(court_name, bold=True, center=True, space_after_pt=6)

    if case_number:
        add_paragraph(f"תיק מס' {case_number}", center=True, space_after_pt=2)
    if case_title:
        add_paragraph(case_title, center=True, space_after_pt=6)

    # Divider line
    p_div = doc.add_paragraph()
    pPr = p_div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a3a5c')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p_div.paragraph_format.space_after = Pt(8)

    add_paragraph("סיכום AI", bold=True, center=True, space_before_pt=6, space_after_pt=10)

    # ── Body text ────────────────────────────────────────────────────────────────
    # Split by blank lines into paragraphs
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]

    for para in paragraphs:
        # Detect section headers: short lines (< 60 chars) not ending with period/comma
        is_header = (len(para) < 60 and
                     not para.endswith(('.', ',', ':', ')', '"', "'")) and
                     not para[0].isdigit())
        if is_header:
            add_paragraph(para, bold=True, space_before_pt=8, space_after_pt=4)
        else:
            add_paragraph(para, space_before_pt=3, space_after_pt=3)

    # ── Footer ───────────────────────────────────────────────────────────────────
    p_div2 = doc.add_paragraph()
    pPr2 = p_div2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'aaaaaa')
    pBdr2.append(top)
    pPr2.append(pBdr2)
    p_div2.paragraph_format.space_before = Pt(12)

    now = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    add_paragraph(f"הופק על ידי מערכת שירה AI  |  {now}", font_size=9, center=True, color=(150, 150, 150))

    # ── Save to buffer ────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"סיכום_AI_{case_number or 'תיק'}_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    from flask import send_file
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename
    )


def create_message_docx(text, case_data, court_name):
    """Generate a formal message docx with court header. Returns io.BytesIO."""
    import io
    import datetime
    import docx as _docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _docx.Document()
    section = doc.sections[0]
    section.page_width    = 7560310
    section.page_height   = 10692130
    section.left_margin   = 900430
    section.right_margin  = 1141095
    section.top_margin    = 600000
    section.bottom_margin = 810260

    body_el = doc.element.body
    sectPr = body_el.get_or_add_sectPr()
    sectPr.append(OxmlElement('w:bidi'))

    FONT = 'FrankRuehl'
    SIZE = 14

    def _rtl_para(p, center=False):
        pf = p.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement('w:bidi'))
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center' if center else 'both')
        pPr.append(jc)

    def _add(text_content, bold=False, center=False, size=None):
        p = doc.add_paragraph()
        _rtl_para(p, center)
        run = p.add_run(text_content)
        run.font.name  = FONT
        run.font.size  = Pt(size or SIZE)
        run.font.bold  = bold
        rPr = run._r.get_or_add_rPr()
        rPr.append(OxmlElement('w:rtl'))
        lang = OxmlElement('w:lang')
        lang.set(qn('w:bidi'), 'he-IL')
        rPr.append(lang)

    hebrew_date = datetime.datetime.now().strftime("%d/%m/%Y")
    file_number = case_data.get("fileNumber", "")
    side_a      = case_data.get("sideA", "")
    side_b      = case_data.get("sideB", "")
    subject     = case_data.get("subject", "")

    _add("בבית הדין הרבני האזורי", bold=True, center=True)
    _add(court_name, bold=True, center=True)
    _add("")
    if file_number:
        _add(f"תיק מס' {file_number}", center=True)
    if side_a or side_b:
        _add(f"{side_a} — {side_b}", center=True)
    if subject:
        _add(f"נושא: {subject}", center=True)
    _add(f"תאריך: {hebrew_date}", center=True)

    # Divider
    p_div = doc.add_paragraph()
    pPr = p_div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '1a3a5c')
    pBdr.append(bot); pPr.append(pBdr)
    p_div.paragraph_format.space_after = Pt(10)

    for line in text.split('\n'):
        _add(line.strip() if line.strip() else "")

    # Footer divider
    p_div2 = doc.add_paragraph()
    pPr2 = p_div2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1');    top.set(qn('w:color'), 'aaaaaa')
    pBdr2.append(top); pPr2.append(pBdr2)
    p_div2.paragraph_format.space_before = Pt(12)

    _add("בית הדין הרבני", bold=True, center=True, size=12)
    _add(court_name, bold=True, center=True, size=12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/api/send-message", methods=["POST"])
def send_message():
    import re as _re, time as _time, getpass as _getpass
    body       = request.json or {}
    text       = body.get("text", "").strip()
    case_data  = body.get("caseData", {})

    if not text:
        return jsonify({"error": "no text provided"}), 400

    file_id    = case_data.get("fileId")
    court_id   = case_data.get("courtId", 5)
    court_name = case_data.get("courtName", "בית הדין הרבני")

    if not file_id:
        return jsonify({"error": "fileId required"}), 400

    # 1. Generate docx
    try:
        buf = create_message_docx(text, case_data, court_name)
    except Exception as e:
        return jsonify({"error": f"docx generation failed: {e}"}), 500

    # 2. Write to UNC Temp path
    try:
        username = _getpass.getuser()
        unc_temp = f"\\\\Prod-nas1\\filer$\\Root\\Data\\Users\\{username}\\ScanDocuments\\Temp"
        filename = f"shiramsg_{file_id}_{int(_time.time())}.docx"
        unc_path = os.path.join(unc_temp, filename)
        with open(unc_path, "wb") as f:
            f.write(buf.read())
        print(f"[send-message] wrote {filename} to UNC")
    except Exception as e:
        return jsonify({"error": f"UNC write failed: {e}"}), 500

    # 3. Call SPFE ImportDocument — shiraDocId=fileId assigns new DocumentID
    try:
        unc_escaped = unc_path.replace("\\", "\\\\")
        spfe_body   = (
            f"{{'fileUrl':'{unc_escaped}', "
            f"'shiraDocId':'{file_id}', "
            f"'courtId':'{court_id}', "
            f"'isReadOnly':'false'}}"
        )
        r = SESSION.post(
            f"{SPFE}/ShiraDocsMngWS.asmx/ImportDocument",
            data=spfe_body,
            headers={"Content-Type": "application/json"},
            timeout=30
        )
        m = _re.search(r'"d"\s*:\s*(-?\d+)', r.text)
        doc_id = int(m.group(1)) if m else -1
        print(f"[send-message] SPFE response: {r.text[:100]}  docId={doc_id}")
    except Exception as e:
        return jsonify({"error": f"SPFE call failed: {e}"}), 500

    if doc_id <= 0:
        return jsonify({"error": f"SPFE returned {doc_id} — document not created"}), 500

    postal_url = f"{SHIRA}/classic/Forms/Postal/Postal.aspx?DocumentIDs={doc_id}&FileID={file_id}"
    return jsonify({"postalUrl": postal_url, "docId": doc_id})


@app.route("/api/send-message-direct", methods=["POST"])
def send_message_direct():
    """
    Option B — send a court email WITHOUT creating any Shira document.
    Uses the open internal SMTP relay (mail.rbc.gov.il:25, no auth).
    """
    import smtplib, html as _html
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    from email.utils import formatdate, make_msgid
    import datetime as _dt

    body       = request.json or {}
    text       = body.get("text", "").strip()
    to_email   = body.get("toEmail", "").strip()
    subject    = body.get("subject", "").strip()
    case_data  = body.get("caseData", {})

    if not text:
        return jsonify({"ok": False, "error": "no text provided"}), 400
    if not to_email or "@" not in to_email:
        return jsonify({"ok": False, "error": "valid toEmail required"}), 400

    court_id   = case_data.get("courtId", 5)
    court_name = case_data.get("courtName") or {
        1:"ירושלים",2:"תל אביב",3:"חיפה",4:"פתח תקוה",5:"רחובות",
        6:"באר שבע",7:"טבריה",8:"צפת",9:"אשדוד",10:"אשקלון",
        11:"נתניה",12:"בית הדין הגדול",13:"אריאל",
    }.get(court_id, f"בית הדין #{court_id}")

    if not subject:
        fn = case_data.get("fileNumber", "")
        subject = f"הודעה מבית הדין הרבני — תיק {fn}" if fn else "הודעה מבית הדין הרבני"

    file_number = case_data.get("fileNumber", "")
    side_a      = case_data.get("sideA", "")
    side_b      = case_data.get("sideB", "")
    subj_case   = case_data.get("subject", "")
    today       = _dt.date.today().strftime("%d/%m/%Y")

    # Plain text
    plain_lines = [f"בית הדין הרבני האזורי {court_name}", "-"*40]
    if file_number: plain_lines.append(f"תיק מס': {file_number}")
    if side_a or side_b: plain_lines.append(f"{side_a} נ' {side_b}")
    if subj_case: plain_lines.append(f"נושא: {subj_case}")
    plain_lines += [f"תאריך: {today}", "-"*40, "", *text.split("\n"),
                    "", "-"*40, f"בית הדין הרבני האזורי {court_name}", "no-reply@rbc.gov.il"]
    plain_body = "\n".join(plain_lines)

    # HTML
    header2 = ""
    if file_number: header2 += f"<div>תיק מס': {_html.escape(file_number)}</div>"
    if side_a or side_b: header2 += f"<div>{_html.escape(side_a)} נ' {_html.escape(side_b)}</div>"
    if subj_case: header2 += f"<div>נושא: {_html.escape(subj_case)}</div>"
    body_html = "<br>".join(_html.escape(l) for l in text.split("\n"))
    html_body = f"""<!DOCTYPE html>
<html dir="rtl" lang="he"><head><meta charset="utf-8"></head>
<body style="font-family:Arial,sans-serif;font-size:14px;color:#222;direction:rtl;">
<table width="600" cellpadding="0" cellspacing="0" style="margin:20px auto;border:1px solid #ccc;">
<tr><td style="background:#1a3a5c;color:#fff;padding:16px;text-align:center;">
  <div style="font-size:18px;font-weight:bold;">בית הדין הרבני האזורי</div>
  <div style="font-size:16px;">{_html.escape(court_name)}</div></td></tr>
<tr><td style="padding:12px 20px;background:#f5f7fa;border-bottom:1px solid #ddd;font-size:13px;color:#555;">
  {header2}<div>תאריך: {today}</div></td></tr>
<tr><td style="padding:20px;line-height:1.7;">{body_html}</td></tr>
<tr><td style="padding:16px;text-align:center;background:#f5f7fa;border-top:1px solid #ddd;font-size:12px;color:#888;">
  בית הדין הרבני האזורי {_html.escape(court_name)} &nbsp;|&nbsp; no-reply@rbc.gov.il</td></tr>
</table></body></html>"""

    msg = MIMEMultipart("mixed")
    msg["Subject"]    = subject
    msg["From"]       = f"בית הדין הרבני האזורי {court_name} <no-reply@rbc.gov.il>"
    msg["To"]         = to_email
    msg["Date"]       = formatdate(localtime=True)
    msg["Message-ID"] = make_msgid(domain="rbc.gov.il")
    msg["X-Mailer"]   = "ShiraAI"
    alt = MIMEMultipart("alternative")
    alt.attach(MIMEText(plain_body, "plain", "utf-8"))
    alt.attach(MIMEText(html_body,  "html",  "utf-8"))
    msg.attach(alt)

    try:
        with smtplib.SMTP("mail.rbc.gov.il", 25, timeout=15) as srv:
            srv.ehlo("rbc.gov.il")
            srv.sendmail("no-reply@rbc.gov.il", [to_email], msg.as_bytes())
        print(f"[send-direct] sent to {to_email}  id={msg['Message-ID']}")
        return jsonify({"ok": True, "messageId": msg["Message-ID"]})
    except Exception as e:
        print(f"[send-direct] SMTP error: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/usage")
def usage_stats():
    """Return aggregated usage stats from the log file."""
    log_path = os.path.join(os.path.dirname(__file__), "usage_log.jsonl")
    if not os.path.exists(log_path):
        return jsonify({"queries": 0, "total_tokens": 0, "total_usd": 0, "total_ils": 0, "entries": []})
    entries = []
    try:
        with open(log_path, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line:
                    entries.append(json.loads(line))
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    total_usd     = sum(e.get("usd", 0)      for e in entries)
    total_ils     = sum(e.get("ils", 0)      for e in entries)
    total_tokens  = sum(e.get("total", 0)    for e in entries)
    total_input   = sum(e.get("input", 0)    for e in entries)
    total_output  = sum(e.get("output", 0)   for e in entries)
    total_think   = sum(e.get("thinking", 0) for e in entries)

    return jsonify({
        "queries":        len(entries),
        "total_tokens":   total_tokens,
        "total_input":    total_input,
        "total_output":   total_output,
        "total_thinking": total_think,
        "total_usd":      round(total_usd, 4),
        "total_ils":      round(total_ils, 3),
        "avg_usd":        round(total_usd / len(entries), 4) if entries else 0,
        "last_queries":   entries[-10:]  # last 10 queries
    })


@app.route("/api/health")
def health():
    return jsonify({"status": "ok", "shira": SHIRA})


if __name__ == "__main__":
    import threading
    import webbrowser
    import sys
    import time

    PORT = 5050
    URL  = f"http://localhost:{PORT}"

    def open_browser():
        time.sleep(1.5)   # wait for Flask to start
        webbrowser.open(URL)

    threading.Thread(target=open_browser, daemon=True).start()

    print(f"ShiraAI running at {URL}")
    app.run(host="127.0.0.1", port=PORT, debug=False, use_reloader=False)
