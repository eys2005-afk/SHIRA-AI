"""
protocol_server.py — Standalone protocol transcription server for rabbinical court hearings.
Runs independently on port 5051. Has zero dependency on shira_proxy.py.

Requirements (same packages already installed for shira_proxy):
    pip install flask flask-cors requests python-docx
"""

import os
import json
import base64
import io
import datetime
import requests

# ── Kill any existing process on port 5051 before starting ───────────────────
def _kill_port(port=5051):
    import subprocess, sys
    try:
        if sys.platform == "win32":
            result = subprocess.check_output(
                f'netstat -ano | findstr :{port}', shell=True
            ).decode(errors='ignore')
            pids = set()
            for line in result.splitlines():
                parts = line.split()
                if len(parts) >= 5 and 'LISTENING' in line:
                    pids.add(parts[-1])
            for pid in pids:
                try:
                    subprocess.call(f'taskkill /PID {pid} /F', shell=True,
                                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    print(f'[startup] closed old process on port {port}, PID {pid}')
                except Exception:
                    pass
    except Exception:
        pass

_kill_port(5051)

from flask import Flask, request, jsonify, Response
from flask_cors import CORS

app = Flask(__name__)
app.config['JSON_AS_ASCII'] = False
CORS(app)

# ── Gemini API key — same key used by shira_proxy ────────────────────────────
GEMINI_API_KEY = "AIzaSyCgutrB9sRoyQHC5mY11LiHWF505VQVD44"

import ssl, urllib3
urllib3.disable_warnings()
ssl._create_default_https_context = ssl._create_unverified_context

# ─────────────────────────────────────────────────────────────────────────────
# Frontend HTML
# ─────────────────────────────────────────────────────────────────────────────
_HTML = """<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>פרוטוקול דיון — שירה AI</title>
<style>
* { box-sizing: border-box; margin: 0; padding: 0; }
body { font-family: 'Segoe UI', Arial, sans-serif; direction: rtl; background: #f4f5f7; color: #1a1a2e; font-size: 14px; }

header {
  background: #1a3a5c; color: #fff;
  padding: 12px 24px;
  display: flex; align-items: center; gap: 12px;
}
header h1 { font-size: 17px; font-weight: 500; }
header .sub { font-size: 12px; opacity: 0.7; margin-top: 2px; }

.container { max-width: 900px; margin: 0 auto; padding: 20px 16px; }

.card {
  background: #fff; border: 1px solid #e0e4ea;
  border-radius: 10px; padding: 20px; margin-bottom: 16px;
}
.card-title {
  font-size: 13px; font-weight: 600; color: #555;
  margin-bottom: 14px; text-transform: uppercase;
  letter-spacing: 0.5px; display: flex; align-items: center; gap: 8px;
}

.row { display: flex; gap: 8px; align-items: center; margin-bottom: 10px; flex-wrap: wrap; }
label { font-size: 12px; color: #666; min-width: 90px; }

input[type=text] {
  flex: 1; min-width: 160px; height: 36px;
  border: 1px solid #d0d5dd; border-radius: 7px;
  padding: 0 10px; font-size: 13px; direction: rtl;
  outline: none; transition: border 0.2s;
}
input[type=text]:focus { border-color: #1a3a5c; box-shadow: 0 0 0 3px rgba(26,58,92,0.1); }

button {
  height: 38px; padding: 0 18px;
  border: 1px solid #d0d5dd; border-radius: 7px;
  background: #fff; color: #1a1a2e;
  font-size: 14px; cursor: pointer;
  transition: all 0.15s; white-space: nowrap;
}
button:hover { background: #f0f4f8; }
button.primary { background: #1a3a5c; color: #fff; border-color: #1a3a5c; }
button.primary:hover { background: #14304d; }
button.success { background: #2e7d32; color: #fff; border-color: #2e7d32; }
button.success:hover { background: #1b5e20; }
button:disabled { opacity: 0.45; cursor: not-allowed; }

/* Recording button */
.rec-btn {
  height: 60px; width: 60px; border-radius: 50%;
  font-size: 24px; padding: 0; border: 3px solid #c62828;
  background: #fff; color: #c62828; flex-shrink: 0;
  display: flex; align-items: center; justify-content: center;
  transition: all 0.2s; cursor: pointer;
}
.rec-btn.recording {
  background: #c62828; color: #fff;
  animation: pulse 1.5s infinite;
}
@keyframes pulse {
  0%,100% { box-shadow: 0 0 0 0 rgba(198,40,40,0.4); }
  50%      { box-shadow: 0 0 0 12px rgba(198,40,40,0); }
}

.rec-status {
  display: flex; align-items: center; gap: 16px;
  padding: 16px; background: #f8f9fb; border-radius: 10px;
}
.rec-state { font-size: 15px; font-weight: 600; color: #1a3a5c; }
.rec-timer {
  font-size: 13px; color: #888; margin-top: 3px;
  font-variant-numeric: tabular-nums; font-family: monospace;
}

/* Transcript */
.transcript-box {
  background: #f8fafc; border: 1px solid #e0e4ea;
  border-radius: 8px; min-height: 340px; max-height: 560px;
  overflow-y: auto; padding: 16px;
  font-size: 14px; line-height: 2.1; direction: rtl;
}
.transcript-box:empty::before {
  content: 'הפרוטוקול יופיע כאן תוך כדי ההקלטה...';
  color: #bbb; font-style: italic;
}

.protocol-line { display: flex; gap: 10px; align-items: flex-start; margin-bottom: 4px; }
.speaker-tag {
  font-weight: 700; min-width: 120px; flex-shrink: 0;
  font-size: 13px; padding-top: 2px;
}
.speaker-tag.judge        { color: #4a148c; }
.speaker-tag.husband      { color: #1565c0; }
.speaker-tag.wife         { color: #880e4f; }
.speaker-tag.atty_h       { color: #0277bd; }
.speaker-tag.atty_w       { color: #ad1457; }
.speaker-tag.other        { color: #2e7d32; }
.speaker-tag.interruption { color: #e65100; font-style: italic; }

.line-text { flex: 1; }
.line-text[contenteditable=true]:focus {
  outline: 2px solid #1a3a5c; border-radius: 4px; padding: 2px 6px;
  background: #fff;
}

.chunk-divider { border: none; border-top: 1px dashed #d0d5dd; margin: 10px 0; }

.processing-line {
  color: #888; font-style: italic;
  display: flex; align-items: center; gap: 8px; padding: 4px 0;
}
.spinner {
  display: inline-block; width: 14px; height: 14px;
  border: 2px solid #ddd; border-top-color: #1a3a5c;
  border-radius: 50%; animation: spin 0.7s linear infinite;
  flex-shrink: 0;
}
@keyframes spin { to { transform: rotate(360deg); } }

.error-line { color: #c62828; font-size: 12px; font-style: italic; padding: 2px 0; }

/* Speaker setup grid */
.speaker-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 10px; }
.speaker-row  { display: flex; align-items: center; gap: 8px; }
.speaker-label {
  font-size: 12px; font-weight: 600; min-width: 100px;
  padding: 4px 8px; border-radius: 6px; text-align: center; flex-shrink: 0;
}
.sl-judge   { background: #ede7f6; color: #4a148c; }
.sl-husband { background: #e3f2fd; color: #1565c0; }
.sl-wife    { background: #fce4ec; color: #880e4f; }
.sl-atty-h  { background: #e1f5fe; color: #0277bd; }
.sl-atty-w  { background: #fce4ec; color: #ad1457; }
.sl-other   { background: #e8f5e9; color: #2e7d32; }

.toolbar {
  display: flex; gap: 8px; flex-wrap: wrap;
  align-items: center; margin-bottom: 10px;
}
#edit-mode-btn.active { background: #fff8e1; border-color: #ffc107; color: #e65100; }

.alert { border-radius: 8px; padding: 10px 14px; font-size: 13px; margin-top: 12px; }
.alert-info { background: #e3f2fd; border: 1px solid #90caf9; color: #0d47a1; }
.alert-warn { background: #fff8e1; border: 1px solid #ffe082; color: #e65100; }
</style>
</head>
<body>

<header>
  <div>
    <h1>📝 פרוטוקול דיון</h1>
    <div class="sub">תמלול אוטומטי עם ייחוס דוברים — שירה AI</div>
  </div>
</header>

<div class="container">

  <!-- ── Setup card ─────────────────────────────────────────────────────── -->
  <div class="card" id="setup-card">
    <div class="card-title">⚙️ פרטי הדיון</div>

    <div class="row">
      <label>מספר תיק:</label>
      <input type="text" id="case-number" placeholder="לדוגמה: 1488524/1" style="max-width:200px;" />
      <label>תאריך:</label>
      <input type="text" id="case-date" style="max-width:130px;" />
    </div>

    <div class="card-title" style="margin-top:18px;">👥 שמות הדוברים</div>
    <div class="speaker-grid">
      <div class="speaker-row">
        <span class="speaker-label sl-judge">אב"ד / דיין</span>
        <input type="text" id="sp-judge" placeholder='כינוי הרכב / שם הדיין' />
      </div>
      <div class="speaker-row">
        <span class="speaker-label sl-husband">הבעל</span>
        <input type="text" id="sp-husband" placeholder="שם הבעל" />
      </div>
      <div class="speaker-row">
        <span class="speaker-label sl-wife">האשה</span>
        <input type="text" id="sp-wife" placeholder="שם האשה" />
      </div>
      <div class="speaker-row">
        <span class="speaker-label sl-atty-h">ב"כ הבעל</span>
        <input type="text" id="sp-atty-h" placeholder='עו"ד מטעם הבעל' />
      </div>
      <div class="speaker-row">
        <span class="speaker-label sl-atty-w">ב"כ האשה</span>
        <input type="text" id="sp-atty-w" placeholder='עו"ד מטעם האשה' />
      </div>
      <div class="speaker-row">
        <span class="speaker-label sl-other">אחר</span>
        <input type="text" id="sp-other" placeholder="עד / מומחה / שם אחר" />
      </div>
    </div>

    <div style="margin-top:18px; display:flex; gap:10px; align-items:center;">
      <button class="primary" onclick="startSession()">▶ פתח דיון</button>
      <span style="font-size:12px;color:#888;">הדפדפן יבקש אישור גישה למיקרופון</span>
    </div>
  </div>

  <!-- ── Recording card ────────────────────────────────────────────────── -->
  <div class="card" id="recording-card" style="display:none;">
    <div class="card-title">🎙 הקלטה</div>

    <div class="rec-status">
      <button class="rec-btn" id="rec-btn" onclick="toggleRecording()">⏺</button>
      <div style="flex:1;">
        <div class="rec-state" id="rec-state">מוכן — לחץ ⏺ להתחלה</div>
        <div class="rec-timer"  id="rec-timer">00:00:00</div>
      </div>
      <div style="display:flex;flex-direction:column;gap:6px;align-items:center;">
        <button onclick="sendCurrentChunk()" id="send-btn"
                style="font-size:12px;height:32px;" disabled>
          ⬆ שלח קטע עכשיו
        </button>
        <span id="chunk-status" style="font-size:11px;color:#888;text-align:center;min-height:16px;"></span>
      </div>
    </div>

    <div id="rec-alert" class="alert alert-info" style="display:none;"></div>

    <!-- Transcript -->
    <div style="margin-top:18px;">
      <div class="toolbar">
        <strong style="font-size:13px;color:#1a3a5c;">📄 פרוטוקול</strong>
        <button id="edit-mode-btn" onclick="toggleEditMode()" style="font-size:12px;height:32px;">
          ✏️ מצב עריכה
        </button>
        <button onclick="clearTranscript()" style="font-size:12px;height:32px;color:#888;">
          🗑 נקה
        </button>
        <div style="margin-right:auto;">
          <button class="success" onclick="exportProtocol()" style="font-size:12px;height:32px;">
            📄 ייצא Word
          </button>
        </div>
      </div>
      <div class="transcript-box" id="transcript-box"></div>
    </div>
  </div>

</div><!-- /container -->

<script>
const PROXY = "http://localhost:5051";

let mediaRecorder   = null;
let audioChunks     = [];
let timerInterval   = null;
let autoSendTimer   = null;
let elapsedSec      = 0;
let isRecording     = false;
let editMode        = false;
const AUTO_SEND_MS  = 45000;   // auto-send every 45 seconds

// Pre-fill today's date
document.getElementById('case-date').value =
  new Date().toLocaleDateString('he-IL', {day:'2-digit',month:'2-digit',year:'numeric'})
            .replace(/\\./g, '/');

// ── Speaker helpers ───────────────────────────────────────────────────────────
function getSpeakers() {
  return {
    judge:   document.getElementById('sp-judge').value.trim()  || 'ביה"ד',
    husband: document.getElementById('sp-husband').value.trim() || 'הבעל',
    wife:    document.getElementById('sp-wife').value.trim()   || 'האשה',
    attyH:   document.getElementById('sp-atty-h').value.trim() || 'ב"כ הבעל',
    attyW:   document.getElementById('sp-atty-w').value.trim() || 'ב"כ האשה',
    other:   document.getElementById('sp-other').value.trim()  || '',
  };
}

// ── Session start ─────────────────────────────────────────────────────────────
async function startSession() {
  try {
    const stream = await navigator.mediaDevices.getUserMedia({ audio: true });
    stream.getTracks().forEach(t => t.stop());
  } catch(e) {
    alert('לא ניתן לגשת למיקרופון:\\n' + e.message);
    return;
  }
  document.getElementById('setup-card').style.display    = 'none';
  document.getElementById('recording-card').style.display = 'block';
}

// ── Recording toggle ──────────────────────────────────────────────────────────
async function toggleRecording() {
  if (!isRecording) { await startRecording(); }
  else               { stopRecording(); }
}

async function startRecording() {
  let stream;
  try {
    stream = await navigator.mediaDevices.getUserMedia({ audio: true });
  } catch(e) {
    alert('שגיאה בפתיחת מיקרופון:\\n' + e.message);
    return;
  }

  const mimeType =
    MediaRecorder.isTypeSupported('audio/webm;codecs=opus') ? 'audio/webm;codecs=opus' :
    MediaRecorder.isTypeSupported('audio/webm')             ? 'audio/webm'             :
                                                              'audio/ogg';
  mediaRecorder = new MediaRecorder(stream, { mimeType });
  audioChunks   = [];

  mediaRecorder.ondataavailable = e => {
    if (e.data && e.data.size > 0) audioChunks.push(e.data);
  };
  mediaRecorder.start(1000);
  isRecording = true;

  document.getElementById('rec-btn').classList.add('recording');
  document.getElementById('rec-btn').textContent  = '⏹';
  document.getElementById('rec-state').textContent = 'מקליט...';
  document.getElementById('send-btn').disabled     = false;
  showAlert('מקליט. הפרוטוקול יתעדכן אוטומטית כל 45 שניות.', 'info');

  autoSendTimer = setInterval(() => sendCurrentChunk(), AUTO_SEND_MS);
  elapsedSec    = 0;
  timerInterval = setInterval(() => {
    elapsedSec++;
    document.getElementById('rec-timer').textContent = fmtTime(elapsedSec);
  }, 1000);
}

function stopRecording() {
  if (mediaRecorder && mediaRecorder.state !== 'inactive') {
    mediaRecorder.stop();
    mediaRecorder.stream.getTracks().forEach(t => t.stop());
  }
  isRecording = false;
  clearInterval(autoSendTimer);
  clearInterval(timerInterval);

  document.getElementById('rec-btn').classList.remove('recording');
  document.getElementById('rec-btn').textContent   = '⏺';
  document.getElementById('rec-state').textContent = 'הקלטה הופסקה';
  document.getElementById('send-btn').disabled     = true;

  if (audioChunks.length > 0) {
    setTimeout(() => sendCurrentChunk(true), 600);
  }
}

// ── Send chunk to server ───────────────────────────────────────────────────────
async function sendCurrentChunk(isFinal = false) {
  if (!audioChunks.length) { setStatus('אין אודיו לשליחה'); return; }

  const chunksToSend = [...audioChunks];
  audioChunks = [];

  const mime = mediaRecorder?.mimeType || 'audio/webm';
  const blob  = new Blob(chunksToSend, { type: mime });

  if (blob.size < 4000) {
    setStatus('קטע קצר מדי, ממתין...');
    audioChunks = chunksToSend;
    return;
  }

  setStatus('מעבד...');
  addProcessingLine();

  const fd = new FormData();
  fd.append('audio',    blob, 'chunk.webm');
  fd.append('speakers', JSON.stringify(getSpeakers()));
  fd.append('is_final', isFinal ? '1' : '0');
  fd.append('case_number', document.getElementById('case-number').value.trim());

  try {
    const resp = await fetch(`${PROXY}/api/transcribe`, { method:'POST', body: fd });
    const data = await resp.json();
    removeProcessingLine();

    if (data.error) {
      setStatus('שגיאה: ' + data.error);
      addErrorLine(data.error);
    } else if (data.lines && data.lines.length) {
      data.lines.forEach(addProtocolLine);
      addDivider();
      setStatus(`✓ ${data.lines.length} שורות`);
    } else {
      setStatus('לא זוהה דיבור');
    }
  } catch(e) {
    removeProcessingLine();
    setStatus('שגיאת רשת');
    addErrorLine(e.message);
  }
}

// ── DOM helpers ───────────────────────────────────────────────────────────────
function addProtocolLine(line) {
  const box = document.getElementById('transcript-box');
  const div = document.createElement('div');
  div.className = 'protocol-line';

  const tag = document.createElement('span');
  tag.className   = 'speaker-tag ' + (line.speaker_class || 'other');
  tag.textContent = line.speaker + ':';

  const txt = document.createElement('span');
  txt.className = 'line-text';
  txt.textContent = line.text;
  if (editMode) txt.contentEditable = 'true';

  div.appendChild(tag);
  div.appendChild(txt);
  box.appendChild(div);
  box.scrollTop = box.scrollHeight;
}

function addDivider() {
  const hr = document.createElement('hr');
  hr.className = 'chunk-divider';
  document.getElementById('transcript-box').appendChild(hr);
}

function addProcessingLine() {
  const div = document.createElement('div');
  div.className = 'processing-line';
  div.id = 'processing-ind';
  div.innerHTML = '<span class="spinner"></span> מתמלל...';
  const box = document.getElementById('transcript-box');
  box.appendChild(div);
  box.scrollTop = box.scrollHeight;
}
function removeProcessingLine() { document.getElementById('processing-ind')?.remove(); }

function addErrorLine(msg) {
  const div = document.createElement('div');
  div.className   = 'error-line';
  div.textContent = '⚠ ' + msg;
  document.getElementById('transcript-box').appendChild(div);
}

function toggleEditMode() {
  editMode = !editMode;
  const btn = document.getElementById('edit-mode-btn');
  btn.classList.toggle('active', editMode);
  btn.textContent = editMode ? '✏️ עריכה פעילה' : '✏️ מצב עריכה';
  document.querySelectorAll('.line-text').forEach(el => {
    el.contentEditable = editMode ? 'true' : 'false';
  });
}

function clearTranscript() {
  if (!confirm('למחוק את כל הפרוטוקול?')) return;
  document.getElementById('transcript-box').innerHTML = '';
}

function showAlert(msg, type='info') {
  const el = document.getElementById('rec-alert');
  el.className     = 'alert alert-' + type;
  el.textContent   = msg;
  el.style.display = 'block';
  if (type === 'info') setTimeout(() => { el.style.display = 'none'; }, 7000);
}

function setStatus(msg) { document.getElementById('chunk-status').textContent = msg; }

function fmtTime(s) {
  return [Math.floor(s/3600), Math.floor((s%3600)/60), s%60]
    .map(n => String(n).padStart(2,'0')).join(':');
}

// ── Export to Word ─────────────────────────────────────────────────────────────
async function exportProtocol() {
  const lines = [];
  document.querySelectorAll('.protocol-line').forEach(div => {
    const speaker_class = [...div.querySelector('.speaker-tag').classList]
                           .find(c => c !== 'speaker-tag') || 'other';
    const speaker = div.querySelector('.speaker-tag').textContent.replace(':','').trim();
    const text    = div.querySelector('.line-text').textContent.trim();
    if (text) lines.push({ speaker, speaker_class, text });
  });

  if (!lines.length) { alert('הפרוטוקול ריק.'); return; }

  try {
    const resp = await fetch(`${PROXY}/api/export`, {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({
        lines,
        case_number: document.getElementById('case-number').value.trim(),
        case_date:   document.getElementById('case-date').value.trim(),
        speakers:    getSpeakers()
      })
    });
    if (!resp.ok) { alert('שגיאה בייצוא'); return; }
    const blob  = await resp.blob();
    const url   = URL.createObjectURL(blob);
    const a     = document.createElement('a');
    const cd    = resp.headers.get('Content-Disposition') || '';
    const match = cd.match(/filename\\*?=(?:UTF-8'')?([^;]+)/i);
    a.download  = match ? decodeURIComponent(match[1]) : 'פרוטוקול.docx';
    a.href = url;
    a.click();
    URL.revokeObjectURL(url);
  } catch(e) { alert('שגיאה: ' + e.message); }
}
</script>
</body>
</html>
"""


# ─────────────────────────────────────────────────────────────────────────────
# Routes
# ─────────────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return Response(_HTML, mimetype="text/html; charset=utf-8")


@app.route("/api/health")
def health():
    return jsonify({"status": "ok", "service": "protocol"})


@app.route("/api/transcribe", methods=["POST"])
def transcribe():
    """
    Receives audio blob from MediaRecorder.
    Sends to Gemini with speaker context.
    Returns structured protocol lines.
    """
    audio_file   = request.files.get("audio")
    speakers_raw = request.form.get("speakers", "{}")

    if not audio_file:
        return jsonify({"error": "no audio file received"}), 400

    try:
        speakers = json.loads(speakers_raw)
    except Exception:
        speakers = {}

    judge   = speakers.get("judge",   'ביה"ד')
    husband = speakers.get("husband", "הבעל")
    wife    = speakers.get("wife",    "האשה")
    atty_h  = speakers.get("attyH",   'ב"כ הבעל')
    atty_w  = speakers.get("attyW",   'ב"כ האשה')
    other   = speakers.get("other",   "")

    speakers_desc = (
        f"- {judge}   (ביה\"ד / דיין / אב\"ד)\n"
        f"- {husband} (הבעל)\n"
        f"- {wife}    (האשה)\n"
        f"- {atty_h}  (ב\"כ הבעל)\n"
        f"- {atty_w}  (ב\"כ האשה)"
    )
    if other:
        speakers_desc += f"\n- {other}  (אחר)"

    speaker_name_map = {
        "judge":   judge,
        "husband": husband,
        "wife":    wife,
        "atty_h":  atty_h,
        "atty_w":  atty_w,
        "other":   other or "אחר",
    }

    audio_bytes = audio_file.read()
    audio_b64   = base64.b64encode(audio_bytes).decode()
    mime_type   = audio_file.mimetype or "audio/webm"
    if not mime_type or mime_type in ("application/octet-stream", ""):
        mime_type = "audio/webm"

    prompt = f"""אתה כותב פרוטוקול לבית דין רבני.
תמלל את הדיון הקולי ויחס כל דיבור לדובר המתאים.

הדוברים בדיון:
{speakers_desc}

כללי פורמט — חובה להקפיד:
1. עבור כל משפט, כתוב שורה בפורמט: SPEAKER_KEY|שם הדובר|הטקסט
2. SPEAKER_KEY הוא אחד מ: judge / husband / wife / atty_h / atty_w / other
3. התפרצות / הפרעה: interruption|התפרצות|[מה נאמר]
4. תמלל בעברית בלבד — שמור על הטקסט המדויק, אל תסכם
5. אם אין דיבור ברור — החזר רק שורה ריקה, ללא הסברים

דוגמה:
judge|{judge}|הדיון פתוח. נא להציג עמדות.
husband|{husband}|אני טוען שהסכם הממון אינו תקף.
atty_w|{atty_w}|אדוני, אני מתנגדת.
interruption|התפרצות|[הבעל מדבר בלי לקבל רשות]
"""

    try:
        url = (
            "https://generativelanguage.googleapis.com/v1beta/models/"
            f"gemini-2.5-flash:generateContent?key={GEMINI_API_KEY}"
        )
        payload = {
            "contents": [{
                "parts": [
                    {"inline_data": {"mime_type": mime_type, "data": audio_b64}},
                    {"text": prompt}
                ]
            }],
            "generationConfig": {"temperature": 0.1, "maxOutputTokens": 4096}
        }

        resp = requests.post(
            url, json=payload,
            proxies={"https": None, "http": None},
            verify=False, timeout=120
        )
        resp.raise_for_status()
        data = resp.json()

        raw_text = (
            data.get("candidates", [{}])[0]
                .get("content", {})
                .get("parts", [{}])[0]
                .get("text", "")
        ).strip()

        print(f"[transcribe] Gemini response ({len(raw_text)} chars):\n{raw_text[:400]}")

        speaker_class_map = {
            "judge":        "judge",
            "husband":      "husband",
            "wife":         "wife",
            "atty_h":       "atty_h",
            "atty_w":       "atty_w",
            "other":        "other",
            "interruption": "interruption",
        }

        lines = []
        for raw_line in raw_text.splitlines():
            raw_line = raw_line.strip()
            if not raw_line or "|" not in raw_line:
                continue
            parts = raw_line.split("|", 2)
            if len(parts) < 3:
                continue
            key, name, text = parts[0].strip(), parts[1].strip(), parts[2].strip()
            if not text:
                continue
            lines.append({
                "speaker_key":   key,
                "speaker":       name or speaker_name_map.get(key, key),
                "speaker_class": speaker_class_map.get(key, "other"),
                "text":          text
            })

        return jsonify({"lines": lines})

    except Exception as e:
        print(f"[transcribe] error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/export", methods=["POST"])
def export_docx():
    """Generate a protocol DOCX in FrankRuehl format."""
    import docx as _docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    body        = request.json
    lines       = body.get("lines", [])
    case_number = body.get("case_number", "")
    case_date   = body.get("case_date", "")
    speakers    = body.get("speakers", {})

    doc = _docx.Document()

    section = doc.sections[0]
    section.page_width    = 7560310
    section.page_height   = 10692130
    section.left_margin   = 900430
    section.right_margin  = 1141095
    section.top_margin    = 500000
    section.bottom_margin = 810260

    body_el = doc.element.body
    sectPr  = body_el.get_or_add_sectPr()
    sectPr.append(OxmlElement('w:bidi'))

    FONT = 'FrankRuehl'

    def rtl_para(p, center=False):
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement('w:bidi'))
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center' if center else 'both')
        pPr.append(jc)
        p.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
        )

    def make_run(p, text, bold=False, size=13, color=None, italic=False):
        r = p.add_run(text)
        r.font.name   = FONT
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = RGBColor(*color)
        rPr = r._r.get_or_add_rPr()
        rPr.append(OxmlElement('w:rtl'))
        lang = OxmlElement('w:lang')
        lang.set(qn('w:bidi'), 'he-IL')
        rPr.append(lang)
        return r

    def add_para(text, bold=False, center=False, size=13, color=None,
                 sb=4, sa=4, italic=False):
        p = doc.add_paragraph()
        rtl_para(p, center)
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        make_run(p, text, bold=bold, size=size, color=color, italic=italic)
        return p

    def add_divider(color='1a3a5c', top=False):
        p   = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        side = OxmlElement('w:top' if top else 'w:bottom')
        side.set(qn('w:val'), 'single')
        side.set(qn('w:sz'), '6' if not top else '4')
        side.set(qn('w:space'), '1')
        side.set(qn('w:color'), color)
        pBdr.append(side)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(8 if top else 2)
        p.paragraph_format.space_after  = Pt(2 if top else 8)

    # ── Header ────────────────────────────────────────────────────────────────
    add_para("בית הדין הרבני האזורי", bold=True, center=True, size=16, sb=6, sa=2)
    add_para("פרוטוקול דיון", bold=True, center=True, size=15, sa=4)

    if case_number:
        add_para(f"תיק מס' {case_number}", center=True, size=12, sa=2)
    if case_date:
        add_para(f"תאריך: {case_date}", center=True, size=12, sa=6)

    if speakers.get("husband") or speakers.get("wife"):
        add_para("הצדדים:", bold=True, size=12, sb=6, sa=2)
        if speakers.get("husband"):
            atty = f"  (ב\"כ: {speakers['attyH']})" if speakers.get("attyH") else ""
            add_para(f"הבעל: {speakers['husband']}{atty}", size=12, sa=1)
        if speakers.get("wife"):
            atty = f"  (ב\"כ: {speakers['attyW']})" if speakers.get("attyW") else ""
            add_para(f"האשה: {speakers['wife']}{atty}", size=12, sa=4)

    add_divider()

    # ── Protocol body ─────────────────────────────────────────────────────────
    COLORS = {
        "judge":        (74,  20, 140),
        "husband":      (21, 101, 192),
        "wife":         (136, 14,  79),
        "atty_h":       (2,  119, 189),
        "atty_w":       (173, 20,  87),
        "other":        (46, 125,  50),
        "interruption": (230, 81,   0),
    }

    for line in lines:
        speaker     = line.get("speaker", "")
        text        = line.get("text", "")
        key         = line.get("speaker_class") or line.get("speaker_key") or "other"
        if not text:
            continue

        p = doc.add_paragraph()
        rtl_para(p)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

        color = COLORS.get(key, (30, 30, 30))
        make_run(p, f"{speaker}: ", bold=True, size=13, color=color)
        make_run(p, text, size=13, italic=(key == "interruption"), color=(color if key == "interruption" else None))

    # ── Footer ────────────────────────────────────────────────────────────────
    add_divider(color='aaaaaa', top=True)
    now = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    add_para(f"הופק על ידי מערכת שירה AI  |  {now}", size=9, center=True,
             color=(150, 150, 150), sb=2, sa=2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = (
        f"פרוטוקול_{case_number or 'דיון'}_"
        f"{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    )

    from flask import send_file
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename
    )


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import threading
    import webbrowser
    import time

    PORT = 5051
    URL  = f"http://localhost:{PORT}"

    def open_browser():
        time.sleep(1.5)
        webbrowser.open(URL)

    threading.Thread(target=open_browser, daemon=True).start()

    print(f"Protocol server running at {URL}")
    app.run(host="127.0.0.1", port=PORT, debug=False, use_reloader=False)
