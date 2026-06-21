"""
Phase A — Step 8
================
The real document-creation chain:
  1. Build a minimal docx in memory
  2. POST it as multipart to IframeFromMyComputerDocument.aspx
     → server writes to UNC temp + creates Shira document record
     → response contains the real DocumentID in hdnDocumentID
  3. GET UploadFileToDM.aspx with that DocumentID  → copies into DM
  4. Print Postal.aspx URL

Run:
    python phase_a_step8.py > step8.txt 2>&1
    notepad step8.txt
"""
import os, io, re, time, urllib.parse
import requests
from requests_negotiate_sspi import HttpNegotiateAuth
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

os.environ['NO_PROXY'] = 'shira2,prod-spfe,localhost,127.0.0.1'
os.environ['no_proxy'] = 'shira2,prod-spfe,localhost,127.0.0.1'

SHIRA      = "http://shira2"
FILE_ID    = "2923739"
ENTITY_ID  = "1936401"
COURT_ID   = "5"
COURT_NAME = "רחובות"
UNC_TEMP   = r"\\Prod-nas1\filer$\Root\Data\Users\elchanans\ScanDocuments\Temp"

TEST_CASE = {
    "fileNumber": "1574928/2",
    "sideA":      "ישראל ישראלי",
    "sideB":      "שרה ישראלי",
    "subject":    "גירושין",
}
TEST_MESSAGE = (
    "הנכם מוזמנים להתייצב לדיון שייערך בבית הדין הרבני האזורי רחובות.\n"
    "אי התייצבות עלולה לגרום לחיוב בהוצאות."
)

# ── Session ───────────────────────────────────────────────────────────────────
def make_session():
    s = requests.Session()
    s.auth      = HttpNegotiateAuth()
    s.trust_env = False
    s.proxies   = {}
    s.verify    = False
    return s

# ── RTL docx builder ──────────────────────────────────────────────────────────
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_rtl_para(doc, text, bold=False, size=12, center=False):
    p = doc.add_paragraph()
    set_rtl(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'David'
    return p

def build_docx(case, message_text, court_name):
    doc = Document()
    for section in doc.sections:
        bidi = OxmlElement('w:bidi')
        section._sectPr.append(bidi)
    add_rtl_para(doc, "בבית הדין הרבני האזורי", bold=True, size=14, center=True)
    add_rtl_para(doc, court_name, bold=True, size=16, center=True)
    doc.add_paragraph()
    add_rtl_para(doc, f"תאריך: {date.today().strftime('%d/%m/%Y')}")
    add_rtl_para(doc, f"תיק מס': {case['fileNumber']}")
    add_rtl_para(doc, f"נושא: {case['subject']}")
    add_rtl_para(doc, f"צד א: {case['sideA']}")
    add_rtl_para(doc, f"צד ב: {case['sideB']}")
    doc.add_paragraph()
    add_rtl_para(doc, "─" * 40, center=True)
    doc.add_paragraph()
    for line in message_text.splitlines():
        add_rtl_para(doc, line, size=12)
    doc.add_paragraph()
    add_rtl_para(doc, "─" * 40, center=True)
    add_rtl_para(doc, "בית הדין הרבני", bold=True, center=True)
    add_rtl_para(doc, court_name, bold=True, center=True)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── Helpers ───────────────────────────────────────────────────────────────────
def get_viewstate(session, url):
    r = session.get(url, timeout=20)
    vs  = re.search(r'id="__VIEWSTATE"\s+value="([^"]*)"', r.text)
    ev  = re.search(r'id="__EVENTVALIDATION"\s+value="([^"]*)"', r.text)
    vsg = re.search(r'id="__VIEWSTATEGENERATOR"\s+value="([^"]*)"', r.text)
    return {
        "__VIEWSTATE":          vs.group(1)  if vs  else "",
        "__EVENTVALIDATION":    ev.group(1)  if ev  else "",
        "__VIEWSTATEGENERATOR": vsg.group(1) if vsg else "",
        "_raw": r.text,
    }

def dump_hints(html, needles):
    """Print any lines in the HTML that contain our keywords."""
    for line in html.splitlines():
        for n in needles:
            if n in line:
                print(f"    HINT: {line.strip()[:200]}")
                break

# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    print("=" * 60)
    print("  Phase A Step 8 — IframeFromMyComputerDocument POST")
    print("=" * 60)

    session = make_session()

    # 1. Build docx
    print("\n[1] Building docx...")
    docx_bytes = build_docx(TEST_CASE, TEST_MESSAGE, COURT_NAME)
    filename   = f"shiramsg_{int(time.time())}.docx"
    print(f"    ✅ {len(docx_bytes)} bytes  filename={filename}")

    # 2. Fetch VIEWSTATE from the iframe page
    iframe_url = (f"{SHIRA}/classic/Forms/Documents/Scan/IframeFromMyComputerDocument.aspx"
                  f"?FileID={FILE_ID}&EntityTypeID=1&EntityID={ENTITY_ID}"
                  f"&DocumentID=0&CourtID={COURT_ID}")
    print(f"\n[2] Fetching iframe VIEWSTATE: {iframe_url}")
    vs = get_viewstate(session, iframe_url)
    print(f"    VIEWSTATE len={len(vs['__VIEWSTATE'])}  EV len={len(vs['__EVENTVALIDATION'])}")
    print(f"    Raw page length: {len(vs['_raw'])}")

    # Dump hints from the iframe page
    print("\n    --- Hints from iframe HTML ---")
    dump_hints(vs['_raw'], [
        "hdnDocumentID", "hdnDestinationTempDir", "btnUpload",
        "action=", "enctype", "FormAction", "__FORM_ACTION",
        "function ", "JS_Submit", "SubmitForm",
    ])
    print("    --- End hints ---")

    # 3. POST the file to the iframe endpoint
    print(f"\n[3] POSTing file to iframe endpoint...")
    post_url = f"{SHIRA}/classic/Forms/Documents/Scan/IframeFromMyComputerDocument.aspx"
    form_fields = {
        "__VIEWSTATE":          vs["__VIEWSTATE"],
        "__VIEWSTATEGENERATOR": vs["__VIEWSTATEGENERATOR"],
        "__EVENTVALIDATION":    vs["__EVENTVALIDATION"],
        "__FORM_ACTION":        "UPLOAD",
        "__FORM_SUBMIT_COUNTER": "1",
        "hdnFileID":            FILE_ID,
        "hdnEntityTypeID":      "1",
        "hdnEntityID":          ENTITY_ID,
        "hdnDocumentID":        "0",
        "hdnDestinationTempDir": UNC_TEMP,
        "cboFileSide":          "0",
        "cboScanType":          "1",
        "cboScanSource":        "5",
    }
    files = {
        "filUpload": (filename, io.BytesIO(docx_bytes),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    }
    try:
        r = session.post(post_url, data=form_fields, files=files,
                         headers={"Referer": iframe_url}, timeout=30)
        print(f"    HTTP {r.status_code}  len={len(r.text)}")

        # Look for DocumentID in response
        m = re.search(r'id="hdnDocumentID"\s+value="(\d+)"', r.text)
        doc_id_from_iframe = int(m.group(1)) if m else None
        print(f"    hdnDocumentID in response: {doc_id_from_iframe}")

        # Dump all hints
        print("\n    --- Hints from POST response ---")
        dump_hints(r.text, ["DocumentID", "docId", "hdnDocument", "Error",
                             "שגיאה", "success", "window.parent", "opener"])
        # Print first 2000 chars for inspection
        print(f"\n    --- First 2000 chars ---")
        print(r.text[:2000])
        print("    --- End ---")

    except Exception as e:
        print(f"    ❌ {e}")
        return

    if not doc_id_from_iframe or doc_id_from_iframe == 0:
        print("\n⚠️  No DocumentID returned from iframe POST.")
        print("    This may mean the file input field name is wrong,")
        print("    or we need to try different form field names.")
        print("\n    Trying alternative field name 'FileUpload1'...")

        vs2 = get_viewstate(session, iframe_url)
        form_fields2 = dict(form_fields)
        form_fields2["__VIEWSTATE"]       = vs2["__VIEWSTATE"]
        form_fields2["__EVENTVALIDATION"] = vs2["__EVENTVALIDATION"]
        files2 = {
            "FileUpload1": (filename, io.BytesIO(docx_bytes),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        }
        try:
            r2 = session.post(post_url, data=form_fields2, files=files2,
                              headers={"Referer": iframe_url}, timeout=30)
            print(f"    HTTP {r2.status_code}  len={len(r2.text)}")
            m2 = re.search(r'id="hdnDocumentID"\s+value="(\d+)"', r2.text)
            doc_id_from_iframe = int(m2.group(1)) if m2 else None
            print(f"    hdnDocumentID: {doc_id_from_iframe}")
            print(f"\n    --- First 2000 chars ---")
            print(r2.text[:2000])
        except Exception as e:
            print(f"    ❌ {e}")
        return

    print(f"\n    ✅ Got DocumentID = {doc_id_from_iframe}")

    # 4. Register in DM
    print(f"\n[4] GET UploadFileToDM.aspx  DocumnetId={doc_id_from_iframe}")
    unc_path  = os.path.join(UNC_TEMP, filename)
    enc_path  = urllib.parse.quote(unc_path, safe='')
    enc_name  = urllib.parse.quote(filename, safe='')
    dm_url    = (f"{SHIRA}/classic/Forms/Documents/DM/UploadFileToDM.aspx"
                 f"?SourceFilePath={enc_path}"
                 f"&DocumnetId={doc_id_from_iframe}"
                 f"&DocumnetTypeId=1"
                 f"&DocName={enc_name}")
    try:
        r3 = session.get(dm_url, timeout=30)
        print(f"    HTTP {r3.status_code}")
        print(f"    Response: {r3.text[:300]}")
    except Exception as e:
        print(f"    ❌ {e}")

    # 5. Postal URL
    postal_url = (f"{SHIRA}/classic/Forms/Postal/Postal.aspx"
                  f"?DocumentIDs={doc_id_from_iframe}&FileID={FILE_ID}")
    print(f"\n[5] Postal URL:")
    print(f"    {postal_url}")
    print("\nOpen that URL in your browser, select recipients, and send.")
    print("=" * 60)

if __name__ == "__main__":
    main()
