"""
phase_b_register_request.py — Register an INCOMING request in a Shira case.
==========================================================================
Phase B: a litigant sends a message (via WhatsApp). ShiraAI turns it into a
formal request document INSIDE the case — no Postal, no manual steps.

This reuses the PROVEN engine from phase_a_playwright.py (which produced a
real DocumentID 14732119), but:
  * the docx is framed as a request FROM the party
  * NO Postal step — once the document exists, the request is registered
  * runs headless by default (for the WhatsApp polling loop); pass
    headless=False to watch it the first time

Public function:
    register_request(file_id, entity_id, court_id, request_text, sender_name,
                     case_meta, scan_type=None, headless=True) -> doc_id | None

Test:
    python phase_b_register_request.py
"""
import os, sys, time, tempfile
from datetime import date

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from playwright.sync_api import sync_playwright

SHIRA      = "http://shira2"

# ── Test parameters (a real case) ───────────────────────────────────────────
FILE_ID    = "2923739"
ENTITY_ID  = "1936401"
COURT_ID   = "5"
COURT_NAME = "רחובות"

TEST_CASE = {
    "fileNumber": "1574928/2",
    "subject":    "גירושין",
    "courtName":  COURT_NAME,
}
TEST_SENDER  = "שמריה אלחנן ישראל"
TEST_REQUEST = (
    "לכבוד בית הדין הרבני,\n"
    "הריני מבקש לדחות את מועד הדיון הקבוע מאחר שאני מאושפז.\n"
    "מצורף אישור רפואי. אודה להיענותכם."
)

# ── RTL docx — framed as an incoming request ────────────────────────────────
def _set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def _para(doc, text, bold=False, size=12, center=False):
    p = doc.add_paragraph(); _set_rtl(p)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'David'
    return p

def build_request_docx(request_text, sender_name, case_meta):
    """A document that reads as a request submitted by a party."""
    doc = Document()
    for s in doc.sections:
        b = OxmlElement('w:bidi'); s._sectPr.append(b)

    _para(doc, "בבית הדין הרבני האזורי", bold=True, size=14, center=True)
    _para(doc, case_meta.get("courtName", ""), bold=True, size=16, center=True)
    doc.add_paragraph()

    _para(doc, "בקשה מטעם בעל דין", bold=True, size=13, center=True)
    doc.add_paragraph()

    _para(doc, f"תאריך: {date.today().strftime('%d/%m/%Y')}")
    _para(doc, f"תיק מס': {case_meta.get('fileNumber','')}")
    _para(doc, f"נושא: {case_meta.get('subject','')}")
    _para(doc, f"מגיש הבקשה: {sender_name}")
    doc.add_paragraph()
    _para(doc, "─" * 40, center=True)
    doc.add_paragraph()

    _para(doc, "תוכן הבקשה:", bold=True)
    for line in request_text.splitlines():
        _para(doc, line, size=12)

    doc.add_paragraph()
    _para(doc, "─" * 40, center=True)
    _para(doc, "הוגש באמצעות מערכת ShiraAI", size=10, center=True)

    path = os.path.join(tempfile.gettempdir(), f"shira_req_{int(time.time())}.docx")
    doc.save(path)
    return path

# ── Engine (proven in phase_a_playwright) ───────────────────────────────────
def _find_file_frame(page):
    for fr in page.frames:
        try:
            if fr.query_selector("#fileUploadMyPcDoc"):
                return fr
        except Exception:
            pass
    ifr = page.query_selector("#ifrDocSource")
    return ifr.content_frame() if ifr else None

def register_request(file_id, entity_id, court_id, request_text, sender_name,
                     case_meta, party_value=None, scan_type=None, headless=True):
    """Create a request document inside the case. Returns real DocumentID or None."""
    docx_path = build_request_docx(request_text, sender_name, case_meta)
    print(f"[doc] {docx_path}")

    upload_url = (f"{SHIRA}/classic/Forms/Documents/Scan/UploadScanDocument.aspx"
                  f"?FileID={file_id}&EntityTypeID=1&EntityID={entity_id}&DocumentID=0")

    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=headless,
            args=["--auth-server-allowlist=*shira2*",
                  "--auth-negotiate-delegate-allowlist=*shira2*"],
        )
        ctx = browser.new_context(ignore_https_errors=True, accept_downloads=True)
        page = ctx.new_page()

        print(f"[1] Opening upload form")
        page.goto(upload_url, wait_until="domcontentloaded")
        page.wait_for_timeout(1500)

        print("[2] Source = My Computer")
        page.check("#rbDocSourceOptionMyComputer")
        page.wait_for_timeout(1500)

        print("[3] Setting file in child iframe")
        frame = _find_file_frame(page)
        if not frame or not frame.query_selector("#fileUploadMyPcDoc"):
            print("   ❌ file input not found"); browser.close(); return None
        frame.set_input_files("#fileUploadMyPcDoc", docx_path)
        page.wait_for_timeout(800)

        print("[4] Party / scan-type / description")
        # party (the litigant who sent the request)
        try:
            opts = page.eval_on_selector(
                "#cboFileSide",
                "el => Array.from(el.options).map(o => ({v:o.value, t:o.text}))")
            print(f"   cboFileSide options: {opts}")
            if party_value:
                page.select_option("#cboFileSide", value=party_value)
            elif len(opts) > 1:
                page.select_option("#cboFileSide", index=1)
            print("   ✅ party selected")
        except Exception as e:
            print(f"   ⚠️ cboFileSide: {e}")

        # scan type — print ALL options so we can pick the right "בקשה" type
        try:
            st_opts = page.eval_on_selector(
                "#cboScanType",
                "el => Array.from(el.options).map(o => ({v:o.value, t:o.text}))")
            print(f"   cboScanType options: {st_opts}")
            if scan_type is not None:
                page.select_option("#cboScanType", value=str(scan_type))
                print(f"   ✅ scanType set to {scan_type}")
        except Exception as e:
            print(f"   ⚠️ cboScanType: {e}")

        try:
            page.fill("#txtDocumentDescription", f"בקשה מאת {sender_name}"[:50])
        except Exception as e:
            print(f"   ⚠️ description: {e}")

        print("[5] Clicking שמור (Save)…")
        page.click("#cmdSave")

        print("[6] Waiting for DocumentID…")
        doc_id = 0
        for _ in range(30):
            page.wait_for_timeout(1000)
            try:
                doc_id = int(page.eval_on_selector("#hdnDocumentID", "el => el.value") or "0")
            except Exception:
                doc_id = 0
            if doc_id > 0:
                break
        print(f"   hdnDocumentID = {doc_id}")

        if doc_id <= 0:
            print("   ❌ no DocumentID")
            if not headless:
                page.wait_for_timeout(20000)
            browser.close()
            return None

        print(f"   ✅ Request registered as DocumentID {doc_id}")
        # let the DM-registration modal finish
        page.wait_for_timeout(4000)

        if not headless:
            print("\n(Visible mode) leaving browser open 20s so you can verify in the case.")
            try:
                page.wait_for_timeout(20000)
            except Exception:
                pass

        browser.close()
        return doc_id

def main():
    print("=" * 60)
    print("  Phase B — register incoming request")
    print("=" * 60)
    doc_id = register_request(
        FILE_ID, ENTITY_ID, COURT_ID,
        TEST_REQUEST, TEST_SENDER, TEST_CASE,
        party_value=None,     # auto-pick first party for the test
        scan_type=None,       # we first PRINT the options, then choose
        headless=False,       # visible for the first test
    )
    print(f"\nResult DocumentID: {doc_id}")
    if doc_id:
        print("✅ Open the case in Shira and confirm the request document appears.")

if __name__ == "__main__":
    main()
