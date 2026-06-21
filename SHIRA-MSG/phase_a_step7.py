"""
Phase A — Step 7: Full automated send-message chain
====================================================
1. Build Hebrew RTL docx in memory
2. Write docx to UNC temp path
3. Fetch VIEWSTATE from UploadScanDocument.aspx
4. POST to UploadScanDocument.aspx → get real Shira DocumentID
5. GET UploadFileToDM.aspx → copies file into DM
6. Print Postal.aspx URL → open in browser to select recipients and send

Run:
    python phase_a_step7.py > step7.txt 2>&1
    notepad step7.txt

Change FILE_ID / ENTITY_ID / COURT_NAME to match a real test case.
"""
import os, io, re, time, urllib.parse
import requests
from requests_negotiate_sspi import HttpNegotiateAuth
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

os.environ['NO_PROXY'] = 'shira2,prod-spfe,localhost,127.0.0.1'
os.environ['no_proxy'] = 'shira2,prod-spfe,localhost,127.0.0.1'

SHIRA      = "http://shira2"
FILE_ID    = "2923739"    # ← change to real case FileID
ENTITY_ID  = "1936401"   # ← change to real case EntityID
COURT_ID   = "5"
COURT_NAME = "רחובות"
UNC_TEMP   = r"\\Prod-nas1\filer$\Root\Data\Users\elchanans\ScanDocuments\Temp"

# The ACTION value confirmed by step 6:
ACTION_SAVE_STAY = "SAVE_STAY"   # ← update if step6 found a different value

TEST_CASE = {
    "fileNumber": "1574928/2",
    "sideA": "ישראל ישראלי",
    "sideB": "שרה ישראלי",
    "subject": "גירושין",
}
TEST_MESSAGE = (
    "הנכם מוזמנים להתייצב לדיון שייערך בבית הדין הרבני האזורי רחובות.\n"
    "אי התייצבות עלולה לגרום לחיוב בהוצאות."
)

# ── Session ───────────────────────────────────────────────────────────────────
def make_session():
    s = requests.Session()
    s.auth      = HttpNegotiateAuth()
    s.trust_env = False
    s.proxies   = {}
    s.verify    = False
    return s

# ── RTL docx builder ──────────────────────────────────────────────────────────
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_rtl_para(doc, text, bold=False, size=12, center=False):
    p = doc.add_paragraph()
    set_rtl(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'David'
    return p

def build_docx(case, message_text, court_name):
    doc = Document()
    for section in doc.sections:
        bidi = OxmlElement('w:bidi')
        section._sectPr.append(bidi)

    add_rtl_para(doc, "בבית הדין הרבני האזורי", bold=True, size=14, center=True)
    add_rtl_para(doc, court_name, bold=True, size=16, center=True)
    doc.add_paragraph()

    today = date.today().strftime('%d/%m/%Y')
    add_rtl_para(doc, f"תאריך: {today}")
    add_rtl_para(doc, f"תיק מס': {case['fileNumber']}")
    add_rtl_para(doc, f"נושא: {case['subject']}")
    add_rtl_para(doc, f"צד א: {case['sideA']}")
    add_rtl_para(doc, f"צד ב: {case['sideB']}")
    doc.add_paragraph()
    add_rtl_para(doc, "─" * 40, center=True)
    doc.add_paragraph()

    for line in message_text.splitlines():
        add_rtl_para(doc, line, size=12)

    doc.add_paragraph()
    add_rtl_para(doc, "─" * 40, center=True)
    doc.add_paragraph()
    add_rtl_para(doc, "בית הדין הרבני", bold=True, center=True)
    add_rtl_para(doc, court_name, bold=True, center=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── Shira helpers ─────────────────────────────────────────────────────────────
def get_viewstate(session, file_id, entity_id):
    url = (f"{SHIRA}/classic/Forms/Documents/Scan/UploadScanDocument.aspx"
           f"?FileID={file_id}&EntityTypeID=1&EntityID={entity_id}&DocumentID=0")
    r = session.get(url, timeout=20)
    vs  = re.search(r'id="__VIEWSTATE"\s+value="([^"]*)"', r.text)
    ev  = re.search(r'id="__EVENTVALIDATION"\s+value="([^"]*)"', r.text)
    vsg = re.search(r'id="__VIEWSTATEGENERATOR"\s+value="([^"]*)"', r.text)
    return {
        "__VIEWSTATE":          vs.group(1)  if vs  else "",
        "__EVENTVALIDATION":    ev.group(1)  if ev  else "",
        "__VIEWSTATEGENERATOR": vsg.group(1) if vsg else "",
    }

def create_shira_document(session, file_id, entity_id, unc_temp):
    """POST to UploadScanDocument.aspx — returns real Shira DocumentID or raises."""
    vs = get_viewstate(session, file_id, entity_id)
    url = f"{SHIRA}/classic/Forms/Documents/Scan/UploadScanDocument.aspx"
    form = {
        "__VIEWSTATE":          vs["__VIEWSTATE"],
        "__VIEWSTATEGENERATOR": vs["__VIEWSTATEGENERATOR"],
        "__EVENTVALIDATION":    vs["__EVENTVALIDATION"],
        "__FORM_ACTION":        ACTION_SAVE_STAY,
        "__FORM_SUBMIT_COUNTER": "1",
        "__SHIRA_USER_ID":      "1438",
        "__SHIRA_COURT_ID":     COURT_ID,
        "__SHIRA_FORMBASE_SCREEN_ID": "47",
        "hdnFileID":            file_id,
        "hdnEntityTypeID":      "1",
        "hdnEntityID":          entity_id,
        "hdnDocumentID":        "0",
        "hdnDestinationTempDir": unc_temp,
        "cboFileSide":          "0",
        "cboScanType":          "1",   # כתב
        "cboScanSource":        "5",   # דוא"ל
    }
    r = session.post(url, data=form, headers={"Referer": url}, timeout=30)
    m = re.search(r'id="hdnDocumentID"\s+value="(\d+)"', r.text)
    if not m or int(m.group(1)) == 0:
        raise RuntimeError(f"UploadScanDocument returned no DocumentID. HTTP {r.status_code}, len={len(r.text)}")
    return int(m.group(1))

def register_in_dm(session, unc_file_path, doc_id, doc_name):
    """GET UploadFileToDM.aspx to copy file from UNC into Shira DM."""
    enc_path = urllib.parse.quote(unc_file_path, safe='')
    enc_name = urllib.parse.quote(doc_name, safe='')
    url = (f"{SHIRA}/classic/Forms/Documents/DM/UploadFileToDM.aspx"
           f"?SourceFilePath={enc_path}"
           f"&DocumnetId={doc_id}"
           f"&DocumnetTypeId=1"    # SCAN
           f"&DocName={enc_name}")
    r = session.get(url, timeout=30)
    return r.status_code, r.text[:500]

# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    print("=" * 60)
    print("  Phase A Step 7 — full automated send-message chain")
    print("=" * 60)

    # 1. Build docx
    print("\n[1] Building Hebrew RTL docx...")
    docx_bytes = build_docx(TEST_CASE, TEST_MESSAGE, COURT_NAME)
    print(f"    ✅ {len(docx_bytes)} bytes")

    # 2. Write to UNC
    ts = int(time.time())
    filename  = f"shiramsg_{ts}.docx"
    unc_path  = os.path.join(UNC_TEMP, filename)
    print(f"\n[2] Writing to UNC: {unc_path}")
    try:
        with open(unc_path, "wb") as f:
            f.write(docx_bytes)
        print("    ✅ Written")
    except Exception as e:
        print(f"    ❌ {e}")
        return

    session = make_session()

    # 3. Create Shira document record → real DocumentID
    print(f"\n[3] Creating Shira document record (UploadScanDocument.aspx)...")
    try:
        doc_id = create_shira_document(session, FILE_ID, ENTITY_ID, UNC_TEMP)
        print(f"    ✅ DocumentID = {doc_id}")
    except Exception as e:
        print(f"    ❌ {e}")
        return

    # 4. Register in DM (UploadFileToDM.aspx)
    print(f"\n[4] Registering in DM (UploadFileToDM.aspx)...")
    try:
        status, snippet = register_in_dm(session, unc_path, doc_id, filename)
        print(f"    HTTP {status}")
        print(f"    Response snippet: {snippet!r}")
        if status == 200:
            print("    ✅ DM registration done")
        else:
            print("    ⚠️  Unexpected status — check Postal.aspx anyway")
    except Exception as e:
        print(f"    ❌ {e}")

    # 5. Postal URL
    postal_url = (f"{SHIRA}/classic/Forms/Postal/Postal.aspx"
                  f"?DocumentIDs={doc_id}&FileID={FILE_ID}")
    print(f"\n[5] Postal URL — open in browser:")
    print(f"    {postal_url}")
    print(f"\n    Select recipients and click Send.\n")
    print("=" * 60)
    print("  Done ✅")
    print("=" * 60)

if __name__ == "__main__":
    main()
