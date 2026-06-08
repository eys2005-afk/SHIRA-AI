"""
shira_msg_test.py
=================
End-to-end test of the send-message chain:
  1. Build a real Hebrew RTL docx with court envelope
  2. Write to UNC temp path
  3. Call SPFE ImportDocument  →  get DocumentID
  4. Print Postal.aspx URL (open manually to verify document appears)

Run on the court server:
    python shira_msg_test.py
"""

import os, io, time, re
import requests
from requests_negotiate_sspi import HttpNegotiateAuth
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Config ────────────────────────────────────────────────────────────────────
SHIRA    = "http://shira2"
SPFE     = "http://prod-spfe:1000"
FILE_ID  = "2923739"          # a real FileID for testing
COURT_ID = "5"
COURT_NAME = "רחובות"
UNC_TEMP = r"\\Prod-nas1\filer$\Root\Data\Users\elchanans\ScanDocuments\Temp"

TEST_MESSAGE = "הנכם מוזמנים להתייצב לדיון שייערך בבית הדין הרבני האזורי רחובות.\nאי התייצבות עלולה לגרום לחיוב בהוצאות."

TEST_CASE = {
    "fileNumber": "1574928/2",
    "sideA": "ישראל ישראלי",
    "sideB": "שרה ישראלי",
    "subject": "גירושין",
}

os.environ['NO_PROXY'] = 'shira2,prod-spfe,localhost,127.0.0.1'
os.environ['no_proxy'] = 'shira2,prod-spfe,localhost,127.0.0.1'

# ── Session ───────────────────────────────────────────────────────────────────
def make_session():
    s = requests.Session()
    s.auth      = HttpNegotiateAuth()
    s.trust_env = False
    s.proxies   = {}   # session-level only — never per-request
    s.verify    = False
    return s

# ── RTL docx builder ──────────────────────────────────────────────────────────
def set_rtl(paragraph):
    """Make a paragraph RTL."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_rtl_para(doc, text, bold=False, size=12, center=False):
    p = doc.add_paragraph()
    set_rtl(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'David'
    return p

def build_docx(case, message_text, court_name):
    doc = Document()

    # Page RTL direction
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        sectPr.append(bidi)

    # Header
    add_rtl_para(doc, "בבית הדין הרבני האזורי", bold=True, size=14, center=True)
    add_rtl_para(doc, court_name, bold=True, size=16, center=True)
    doc.add_paragraph()

    # Envelope
    from datetime import date
    today = date.today().strftime('%d/%m/%Y')
    add_rtl_para(doc, f"תאריך: {today}")
    add_rtl_para(doc, f"תיק מס': {case['fileNumber']}")
    add_rtl_para(doc, f"נושא: {case['subject']}")
    add_rtl_para(doc, f"צד א: {case['sideA']}")
    add_rtl_para(doc, f"צד ב: {case['sideB']}")
    doc.add_paragraph()

    # Separator line (dashes)
    add_rtl_para(doc, "─" * 40, center=True)
    doc.add_paragraph()

    # Message body
    for line in message_text.splitlines():
        add_rtl_para(doc, line, size=12)

    doc.add_paragraph()
    add_rtl_para(doc, "─" * 40, center=True)
    doc.add_paragraph()

    # Footer
    add_rtl_para(doc, "בית הדין הרבני", bold=True, center=True)
    add_rtl_para(doc, court_name, bold=True, center=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── Main test ─────────────────────────────────────────────────────────────────
def main():
    print("=" * 55)
    print("  SHIRA-MSG end-to-end test")
    print("=" * 55)

    # Step 1: build docx
    print("\n[1] Building Hebrew RTL docx...")
    try:
        docx_bytes = build_docx(TEST_CASE, TEST_MESSAGE, COURT_NAME)
        print(f"    ✅ docx built — {len(docx_bytes)} bytes")
    except Exception as e:
        print(f"    ❌ {e}")
        return

    # Step 2: write to UNC
    filename  = f"shiramsg_test_{int(time.time())}.docx"
    unc_path  = os.path.join(UNC_TEMP, filename)
    unc_esc   = unc_path.replace("\\", "\\\\")

    print(f"\n[2] Writing to UNC: {unc_path}")
    try:
        with open(unc_path, "wb") as f:
            f.write(docx_bytes)
        print(f"    ✅ Written")
    except Exception as e:
        print(f"    ❌ {e}")
        return

    # Step 3: SPFE ImportDocument
    print(f"\n[3] Calling SPFE ImportDocument (shiraDocId={FILE_ID})...")
    session = make_session()
    try:
        body = (
            f"{{'fileUrl':'{unc_esc}', "
            f"'shiraDocId':'{FILE_ID}', "
            f"'courtId':'{COURT_ID}', "
            f"'isReadOnly':'false'}}"
        )
        r = session.post(
            f"{SPFE}/ShiraDocsMngWS.asmx/ImportDocument",
            data=body,
            headers={"Content-Type": "application/json"},
            timeout=30,
        )
        print(f"    HTTP {r.status_code}")
        print(f"    Response: {r.text[:300]}")

        m = re.search(r'"d"\s*:\s*(-?\d+)', r.text)
        if not m:
            print("    ❌ Could not parse DocumentID from response")
            return

        doc_id = int(m.group(1))
        if doc_id <= 0:
            print(f"    ❌ Got {doc_id} — unexpected negative/zero result")
            return

        print(f"\n    ✅ DocumentID = {doc_id}")
    except Exception as e:
        print(f"    ❌ {e}")
        return

    # Step 4: print postal URL
    postal_url = f"{SHIRA}/classic/Forms/Postal/Postal.aspx?DocumentIDs={doc_id}&FileID={FILE_ID}"
    print(f"\n[4] Postal URL (open in browser):")
    print(f"    {postal_url}")
    print(f"\n    Open it, check that the document '{filename}' appears,")
    print(f"    select recipients and click Send.\n")

    print("=" * 55)
    print("  TEST PASSED ✅")
    print("=" * 55)

if __name__ == "__main__":
    main()
