"""
phase_a_playwright.py — Full automated send-message via Shira's REAL form.
=========================================================================
Instead of reverse-engineering Shira's cross-iframe async postback, we DRIVE
the real screen with Playwright (bundled Chromium + Windows/NTLM auth):

  1. Build a Hebrew RTL docx, save to a local temp file
  2. Open UploadScanDocument.aspx for the case
  3. Choose "מהמחשב" (My Computer) source
  4. Set the file input (fileUploadMyPcDoc) inside the child iframe
  5. Select a party (cboFileSide), fill description + scan type
  6. Click "שמור" — Shira's own JS does upload + save + DM registration
  7. Read the real DocumentID from hdnDocumentID
  8. Open Postal.aspx with that DocumentID → select recipients & send

Run a test:
    python phase_a_playwright.py

Edit FILE_ID / ENTITY_ID below to a real case before running.
"""
import io, os, sys, time, tempfile
from datetime import date

# ── docx builder ────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from playwright.sync_api import sync_playwright

SHIRA      = "http://shira2"
FILE_ID    = "2923739"
ENTITY_ID  = "1936401"
COURT_ID   = "5"
COURT_NAME = "רחובות"

TEST_CASE = {
    "fileNumber": "1574928/2",
    "sideA":      "ישראל ישראלי",
    "sideB":      "שרה ישראלי",
    "subject":    "גירושין",
}
TEST_MESSAGE = (
    "הנכם מוזמנים להתייצב לדיון שייערך בבית הדין הרבני האזורי רחובות.\n"
    "אי התייצבות עלולה לגרום לחיוב בהוצאות."
)

# ── RTL docx ────────────────────────────────────────────────────────────────
def _set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def _para(doc, text, bold=False, size=12, center=False):
    p = doc.add_paragraph(); _set_rtl(p)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'David'
    return p

def build_docx_file(case, message_text, court_name):
    doc = Document()
    for s in doc.sections:
        b = OxmlElement('w:bidi'); s._sectPr.append(b)
    _para(doc, "בבית הדין הרבני האזורי", bold=True, size=14, center=True)
    _para(doc, court_name, bold=True, size=16, center=True)
    doc.add_paragraph()
    _para(doc, f"תאריך: {date.today().strftime('%d/%m/%Y')}")
    _para(doc, f"תיק מס': {case['fileNumber']}")
    _para(doc, f"נושא: {case['subject']}")
    _para(doc, f"צד א: {case['sideA']}")
    _para(doc, f"צד ב: {case['sideB']}")
    doc.add_paragraph()
    _para(doc, "─" * 40, center=True)
    doc.add_paragraph()
    for line in message_text.splitlines():
        _para(doc, line, size=12)
    doc.add_paragraph()
    _para(doc, "─" * 40, center=True)
    _para(doc, "בית הדין הרבני", bold=True, center=True)
    _para(doc, court_name, bold=True, center=True)

    path = os.path.join(tempfile.gettempdir(), f"shiramsg_{int(time.time())}.docx")
    doc.save(path)
    return path

# ── Playwright driver ───────────────────────────────────────────────────────
def send_message(file_id, entity_id, court_id, docx_path, headless=False):
    upload_url = (f"{SHIRA}/classic/Forms/Documents/Scan/UploadScanDocument.aspx"
                  f"?FileID={file_id}&EntityTypeID=1&EntityID={entity_id}&DocumentID=0")

    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=headless,
            args=[
                "--auth-server-allowlist=*shira2*",
                "--auth-negotiate-delegate-allowlist=*shira2*",
            ],
        )
        ctx = browser.new_context(ignore_https_errors=True, accept_downloads=True)
        page = ctx.new_page()

        # auto-accept any popup/modal windows (DM registration uses a modal)
        ctx.on("page", lambda pop: print(f"   [popup] {pop.url}"))

        print(f"[1] Opening: {upload_url}")
        page.goto(upload_url, wait_until="domcontentloaded")
        page.wait_for_timeout(1500)

        # 2. choose "My Computer" source
        print("[2] Selecting source = My Computer")
        page.check("#rbDocSourceOptionMyComputer")
        page.wait_for_timeout(1500)  # let DocSourceOptionClick swap the child iframe

        # 3. set the file input inside the child iframe
        print("[3] Locating child iframe + file input fileUploadMyPcDoc")
        target_frame = None
        for fr in page.frames:
            try:
                if fr.query_selector("#fileUploadMyPcDoc"):
                    target_frame = fr
                    break
            except Exception:
                pass
        if not target_frame:
            # iframe may need explicit name; try by element handle
            ifr = page.query_selector("#ifrDocSource")
            if ifr:
                target_frame = ifr.content_frame()

        if not target_frame or not target_frame.query_selector("#fileUploadMyPcDoc"):
            print("   ❌ Could not find #fileUploadMyPcDoc in any frame.")
            print("   Frames present:")
            for fr in page.frames:
                print(f"     - {fr.url}")
            page.wait_for_timeout(4000)
            browser.close()
            return None

        target_frame.set_input_files("#fileUploadMyPcDoc", docx_path)
        print(f"   ✅ File set: {docx_path}")
        page.wait_for_timeout(1000)

        # 4. parent fields: party side + description + scan type
        print("[4] Filling parent fields (cboFileSide / description / scanType)")
        # select first real party option (index 0 is the empty placeholder)
        try:
            opts = page.eval_on_selector(
                "#cboFileSide",
                "el => Array.from(el.options).map(o => ({v:o.value, t:o.text}))"
            )
            print(f"   cboFileSide options: {opts[:6]}")
            if len(opts) > 1:
                page.select_option("#cboFileSide", index=1)
                print(f"   ✅ Selected party: {opts[1]}")
        except Exception as e:
            print(f"   ⚠️ cboFileSide: {e}")

        try:
            page.fill("#txtDocumentDescription", "הודעת בית הדין")
        except Exception as e:
            print(f"   ⚠️ description: {e}")

        # 5. click Save (שמור) — Shira's JS does the rest
        print("[5] Clicking שמור (Save)…")
        page.click("#cmdSave")

        # 6. wait for hdnDocumentID to become > 0 (poll up to ~30s)
        print("[6] Waiting for real DocumentID…")
        doc_id = 0
        for _ in range(30):
            page.wait_for_timeout(1000)
            try:
                val = page.eval_on_selector(
                    "#hdnDocumentID", "el => el.value"
                )
                doc_id = int(val or "0")
            except Exception:
                doc_id = 0
            if doc_id > 0:
                break
        print(f"   hdnDocumentID = {doc_id}")

        if doc_id <= 0:
            print("   ❌ No DocumentID yet. Leaving browser open 30s to inspect.")
            page.wait_for_timeout(30000)
            browser.close()
            return None

        print(f"   ✅ DocumentID = {doc_id}")

        # 7. give DM-registration modal time to finish, then open Postal
        page.wait_for_timeout(4000)
        postal_url = (f"{SHIRA}/classic/Forms/Postal/Postal.aspx"
                      f"?DocumentIDs={doc_id}&FileID={file_id}")
        print(f"[7] Opening Postal: {postal_url}")
        page.goto(postal_url, wait_until="domcontentloaded")
        page.wait_for_timeout(3000)

        print("\n✅ DONE. Browser stays open — select recipients & send.")
        print("   (Close the window when finished.)")
        page.wait_for_timeout(600_000)
        browser.close()
        return doc_id

def main():
    print("=" * 60)
    print("  Phase A — Playwright send-message")
    print("=" * 60)
    print("\nBuilding docx…")
    docx_path = build_docx_file(TEST_CASE, TEST_MESSAGE, COURT_NAME)
    print(f"  docx: {docx_path}")
    doc_id = send_message(FILE_ID, ENTITY_ID, COURT_ID, docx_path, headless=False)
    print(f"\nResult DocumentID: {doc_id}")

if __name__ == "__main__":
    main()
